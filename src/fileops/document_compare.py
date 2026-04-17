from __future__ import annotations

import json
import re
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any

from .ai_assistant import generate_compare_ai_report
from .models import OperationResult, OperationStatus
from .utils import duration_ms, ensure_workspace_path, now_iso, unique_path
from .word_template import apply_template_format_to_document

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


REFERENCE_HEADINGS = {"参考文献", "references", "reference"}
TOC_DOT_REGEX = re.compile(r"^(?P<title>.+?)\s*[.·•…]{2,}\s*(?P<page>\d+)\s*$")
EQUATION_LABEL_REGEX = re.compile(
    r"(?P<left>[（(])\s*(?P<chapter>\d+)\s*(?P<sep>[-.])\s*(?P<index>\d+)\s*(?P<right>[）)])\s*$"
)
FIGURE_LABEL_REGEX = re.compile(
    r"^(?P<prefix>图|Figure)\s*(?P<chapter>\d+)\s*(?P<sep>[-.])\s*(?P<index>\d+)\b",
    flags=re.IGNORECASE,
)

CATEGORY_LABELS = {
    "style_mismatch": "样式不一致",
    "direct_format_override": "段落格式不一致",
    "equation_numbering": "公式编号异常",
    "figure_numbering": "图编号异常",
    "reference_format": "参考文献格式异常",
}

CATEGORY_SEVERITY = {
    "style_mismatch": "medium",
    "direct_format_override": "medium",
    "equation_numbering": "high",
    "figure_numbering": "high",
    "reference_format": "medium",
}

SEVERITY_LABELS = {"high": "高", "medium": "中", "low": "低"}


def compare_documents_with_template(
    sources: list[Path],
    destination: Path,
    workspace: Path,
    dry_run: bool,
    template_path: Path,
    detailed_report_path: Path | None = None,
    ai_assist_config: dict[str, Any] | None = None,
) -> list[OperationResult]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    destination = destination.resolve(strict=False)
    ensure_workspace_path(destination, workspace)
    if not dry_run:
        destination.mkdir(parents=True, exist_ok=True)

    template_path = template_path.resolve(strict=False)
    if not template_path.exists():
        raise FileNotFoundError(f"Template does not exist: {template_path}")
    if template_path.suffix.lower() != ".docx":
        raise ValueError("Template must be a .docx file.")

    template_doc = Document(str(template_path))
    template_rules = _extract_template_rules(template_doc)

    results: list[OperationResult] = []
    for source in sources:
        started = datetime.now()
        started_at = now_iso()
        try:
            ensure_workspace_path(source, workspace)
            if not source.exists():
                raise FileNotFoundError(f"Source does not exist: {source}")
            if source.is_dir():
                raise IsADirectoryError(f"Document compare supports files only: {source}")
            if source.suffix.lower() != ".docx":
                raise ValueError("Document compare supports .docx files only.")

            report_json_path = destination / f"{source.stem}_compare_report.json"
            report_docx_path = destination / f"{source.stem}_compare_report.docx"
            adjusted_path = destination / f"{source.stem}_adjusted.docx"
            ai_report_path = destination / f"{source.stem}_ai_assist.md"
            if report_json_path.exists():
                report_json_path = unique_path(report_json_path)
            if report_docx_path.exists():
                report_docx_path = unique_path(report_docx_path)
            if adjusted_path.exists():
                adjusted_path = unique_path(adjusted_path)
            if ai_report_path.exists():
                ai_report_path = unique_path(ai_report_path)

            analysis = _analyze_document_against_template(
                source,
                template_rules,
                template_name=template_path.name,
            )
            issue_count = int(analysis["summary"]["total_issues"])
            status_text = str(analysis["overview"]["status_text"])

            if dry_run:
                message = f"Would compare with template '{template_path.name}', detected {issue_count} issue(s), status: {status_text}."
                results.append(_build_result("doc_compare", source, destination, OperationStatus.DRY_RUN, message, started, started_at))
                continue

            report_json_path.write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8-sig")
            _write_compare_report_docx(report_docx_path, analysis)
            if detailed_report_path is not None:
                _write_compare_report_by_suffix(detailed_report_path, analysis)

            apply_template_format_to_document(source=source, template=template_path, output=adjusted_path)

            ai_suffix = ""
            if ai_assist_config:
                try:
                    generate_compare_ai_report(
                        analysis=analysis,
                        output_path=ai_report_path,
                        config=ai_assist_config,
                    )
                    ai_suffix = f"; AI建议: {ai_report_path.name}"
                except Exception as exc:  # noqa: BLE001
                    ai_suffix = f"; AI建议生成失败: {exc}"

            message = (
                f"Compared with template '{template_path.name}', detected {issue_count} issue(s), status: {status_text}. "
                f"Adjusted file: {adjusted_path.name}; reports: {report_docx_path.name}, {report_json_path.name}{ai_suffix}"
            )
            results.append(_build_result("doc_compare", source, adjusted_path, OperationStatus.SUCCESS, message, started, started_at))
        except Exception as exc:  # noqa: BLE001
            results.append(_build_result("doc_compare", source, None, OperationStatus.FAILED, str(exc), started, started_at))

    return results


def analyze_document_with_template(source: Path, template_path: Path) -> dict[str, Any]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    source = source.resolve(strict=False)
    template_path = template_path.resolve(strict=False)

    if not source.exists():
        raise FileNotFoundError(f"Source does not exist: {source}")
    if source.is_dir():
        raise IsADirectoryError(f"Document compare supports files only: {source}")
    if source.suffix.lower() != ".docx":
        raise ValueError("Document compare supports .docx files only.")

    if not template_path.exists():
        raise FileNotFoundError(f"Template does not exist: {template_path}")
    if template_path.suffix.lower() != ".docx":
        raise ValueError("Template must be a .docx file.")

    template_doc = Document(str(template_path))
    template_rules = _extract_template_rules(template_doc)
    return _analyze_document_against_template(
        source,
        template_rules,
        template_name=template_path.name,
    )


def _extract_template_rules(template_doc: Any) -> dict[str, Any]:
    normal_style = _resolve_first_style_name(template_doc, ("Normal", "正文", "Body Text")) or "Normal"

    heading_styles: dict[int, str] = {}
    toc_styles: dict[int, str] = {}
    for level in range(1, 7):
        heading_name = _resolve_template_heading_style(template_doc, level)
        if heading_name:
            heading_styles[level] = heading_name
        toc_name = _resolve_template_toc_style(template_doc, level)
        if toc_name:
            toc_styles[level] = toc_name

    profile_map = _build_template_profile_map(template_doc, normal_style, heading_styles, toc_styles)
    numbering_rules = _extract_numbering_rules(template_doc)

    return {
        "normal_style": normal_style,
        "heading_styles": heading_styles,
        "toc_styles": toc_styles,
        "profiles": profile_map,
        "equation_rule": numbering_rules["equation"],
        "figure_rule": numbering_rules["figure"],
    }


def _build_template_profile_map(
    template_doc: Any,
    normal_style: str,
    heading_styles: dict[int, str],
    toc_styles: dict[int, str],
) -> dict[str, Any]:
    profiles: dict[str, Any] = {
        "normal": None,
        "heading": {},
        "toc": {},
        "reference_heading": None,
        "reference_entry": None,
    }
    context: dict[str, bool] = {"toc_mode": False, "in_references": False}
    seen_heading_style = False

    for paragraph in template_doc.paragraphs:
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text:
            continue

        if _is_reference_heading_text(text):
            context["in_references"] = True
            context["toc_mode"] = False
            if profiles["reference_heading"] is None:
                profiles["reference_heading"] = _capture_profile(paragraph)
            continue

        if _is_toc_intro_text(text):
            context["toc_mode"] = True

        if bool(context.get("toc_mode")):
            manual_level = _infer_manual_toc_level_from_text(text)
            if manual_level is not None and manual_level not in profiles["toc"]:
                profiles["toc"][manual_level] = _capture_profile(paragraph)
                continue
            if manual_level is None:
                context["toc_mode"] = False

        if bool(context.get("in_references")) and _is_reference_entry_text(paragraph, text):
            candidate = _capture_profile(paragraph)
            existing = profiles["reference_entry"]
            if existing is None or _is_better_reference_profile(candidate, existing):
                profiles["reference_entry"] = candidate
            continue

        style_heading_level = _heading_level_from_style(paragraph)
        if style_heading_level is not None and style_heading_level not in profiles["heading"]:
            profiles["heading"][style_heading_level] = _capture_profile(paragraph)
            seen_heading_style = True
            continue
        if style_heading_level is not None:
            seen_heading_style = True

        toc_level = _toc_level_from_style(paragraph)
        if toc_level is not None and toc_level not in profiles["toc"]:
            profiles["toc"][toc_level] = _capture_profile(paragraph)
            continue

        if profiles["normal"] is None and seen_heading_style and _is_body_text_candidate(text):
            profiles["normal"] = _capture_profile(paragraph)

    if profiles["normal"] is None:
        for paragraph in template_doc.paragraphs:
            text = str(getattr(paragraph, "text", "") or "").strip()
            if not text:
                continue
            if _is_body_text_candidate(text):
                profiles["normal"] = _capture_profile(paragraph)
                break

    if profiles["normal"] is None:
        profiles["normal"] = {"style": normal_style}
    if not str(profiles["normal"].get("style") or "").strip():
        profiles["normal"]["style"] = normal_style

    for level, style_name in heading_styles.items():
        if level not in profiles["heading"]:
            profiles["heading"][level] = {"style": style_name}
    for level, style_name in toc_styles.items():
        if level not in profiles["toc"]:
            profiles["toc"][level] = {"style": style_name}

    if profiles["reference_heading"] is None:
        profiles["reference_heading"] = profiles["heading"].get(1) or profiles["normal"]
    if profiles["reference_entry"] is None:
        profiles["reference_entry"] = profiles["normal"]
    return profiles


def _extract_numbering_rules(template_doc: Any) -> dict[str, dict[str, Any]]:
    equation_rule: dict[str, Any] = {"separator": "-", "left_paren": "(", "right_paren": ")"}
    figure_rule: dict[str, Any] = {"separator": "-", "prefix": "图"}

    for paragraph in template_doc.paragraphs:
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text:
            continue
        equation = EQUATION_LABEL_REGEX.search(text)
        if equation:
            equation_rule = {
                "separator": equation.group("sep"),
                "left_paren": equation.group("left"),
                "right_paren": equation.group("right"),
            }
            break

    for paragraph in template_doc.paragraphs:
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text:
            continue
        figure = FIGURE_LABEL_REGEX.match(text)
        if figure:
            figure_rule = {
                "separator": figure.group("sep"),
                "prefix": str(figure.group("prefix")),
            }
            break

    return {"equation": equation_rule, "figure": figure_rule}


def _analyze_document_against_template(source: Path, rules: dict[str, Any], template_name: str) -> dict[str, Any]:
    source_doc = Document(str(source))

    issues: list[dict[str, Any]] = []
    equation_labels: list[dict[str, Any]] = []
    figure_labels: list[dict[str, Any]] = []
    context: dict[str, bool] = {"toc_mode": False, "in_references": False}

    for index, paragraph in enumerate(source_doc.paragraphs, start=1):
        text = str(paragraph.text or "")
        stripped = text.strip()
        if not stripped:
            continue

        role, level = _classify_source_paragraph(paragraph, stripped, context)
        expected_profile = _expected_profile_for_role(rules, role, level)
        expected_style = str(expected_profile.get("style") or "")
        actual_style = str(getattr(paragraph.style, "name", "") or "")

        if expected_style and actual_style != expected_style:
            category = "reference_format" if role == "reference_entry" else "style_mismatch"
            issues.append(
                _issue(
                    category=category,
                    paragraph=index,
                    text=stripped,
                    detail="段落样式与模板不一致",
                    expected=expected_style,
                    actual=actual_style or "(none)",
                    adjustment=f"将该段设置为样式“{expected_style}”。",
                )
            )

        format_mismatch = _compare_paragraph_format(paragraph, expected_profile)
        if format_mismatch is not None:
            category = "reference_format" if role == "reference_entry" else "direct_format_override"
            issues.append(
                _issue(
                    category=category,
                    paragraph=index,
                    text=stripped,
                    detail=format_mismatch["detail"],
                    expected=format_mismatch["expected"],
                    actual=format_mismatch["actual"],
                    adjustment=format_mismatch["adjustment"],
                )
            )

        equation = _extract_equation_label(stripped, rules["equation_rule"])
        if equation is not None:
            equation["paragraph"] = index
            equation_labels.append(equation)

        figure = _extract_figure_label(stripped, rules["figure_rule"])
        if figure is not None:
            figure["paragraph"] = index
            figure_labels.append(figure)

    issues.extend(_check_equation_numbering(equation_labels))
    issues.extend(_check_figure_numbering(figure_labels))

    category_counter = Counter(item["category"] for item in issues)
    summary = {
        "total_issues": len(issues),
        "style_mismatch": category_counter.get("style_mismatch", 0),
        "direct_format_override": category_counter.get("direct_format_override", 0),
        "equation_numbering": category_counter.get("equation_numbering", 0),
        "figure_numbering": category_counter.get("figure_numbering", 0),
        "reference_format": category_counter.get("reference_format", 0),
    }

    status = "pass" if summary["total_issues"] == 0 else "needs_fix"
    overview = {
        "report_id": datetime.now().strftime("%Y%m%d%H%M%S"),
        "template_name": template_name,
        "source_name": source.name,
        "checked_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "status": status,
        "status_text": "格式检测合格" if status == "pass" else "存在格式差异，需调整",
        "total_pages": len(source_doc.sections),
    }

    categories = [
        {"category": key, "label": CATEGORY_LABELS[key], "count": int(summary.get(key, 0))}
        for key in ("style_mismatch", "direct_format_override", "equation_numbering", "figure_numbering", "reference_format")
    ]

    return {
        "source": str(source),
        "template": template_name,
        "overview": overview,
        "summary": summary,
        "categories": categories,
        "issues": issues,
    }


def _classify_source_paragraph(paragraph: Any, text: str, context: dict[str, bool]) -> tuple[str, int | None]:
    if _is_reference_heading_text(text):
        context["in_references"] = True
        context["toc_mode"] = False
        return "reference_heading", None

    if _is_toc_intro_text(text):
        context["toc_mode"] = True
        return "normal", None

    if bool(context.get("toc_mode")):
        toc_level = _infer_manual_toc_level_from_text(text)
        if toc_level is not None:
            return "toc", toc_level
        context["toc_mode"] = False

    if bool(context.get("in_references")) and _is_reference_entry_text(paragraph, text):
        return "reference_entry", None

    toc_level = _detect_toc_level(paragraph, text)
    if toc_level is not None:
        return "toc", toc_level

    heading_level = _detect_heading_level(paragraph, text)
    if heading_level is not None:
        return "heading", heading_level

    return "normal", None


def _expected_profile_for_role(rules: dict[str, Any], role: str, level: int | None) -> dict[str, Any]:
    profiles = rules.get("profiles", {})
    normal_profile = dict(profiles.get("normal") or {"style": rules.get("normal_style", "Normal")})
    if role == "reference_heading":
        return dict(profiles.get("reference_heading") or profiles.get("heading", {}).get(1) or normal_profile)
    if role == "reference_entry":
        return dict(profiles.get("reference_entry") or normal_profile)
    if role == "toc":
        toc_map = profiles.get("toc", {})
        if level is not None and level in toc_map:
            return dict(toc_map[level])
        if 1 in toc_map:
            return dict(toc_map[1])
        return normal_profile
    if role == "heading":
        heading_map = profiles.get("heading", {})
        if level is not None and level in heading_map:
            return dict(heading_map[level])
        if 1 in heading_map:
            return dict(heading_map[1])
        return normal_profile
    return normal_profile


def _compare_paragraph_format(paragraph: Any, expected_profile: dict[str, Any]) -> dict[str, str] | None:
    expected = {
        "首行缩进": _safe_int(expected_profile.get("first_line_indent")),
        "左缩进": _safe_int(expected_profile.get("left_indent")),
        "右缩进": _safe_int(expected_profile.get("right_indent")),
        "段前": _safe_int(expected_profile.get("space_before")),
        "段后": _safe_int(expected_profile.get("space_after")),
        "行距": _expected_line_spacing_text(expected_profile),
        "对齐": _alignment_to_text(expected_profile.get("alignment")),
    }
    actual = {
        "首行缩进": _length_to_text(getattr(paragraph.paragraph_format, "first_line_indent", None)),
        "左缩进": _length_to_text(getattr(paragraph.paragraph_format, "left_indent", None)),
        "右缩进": _length_to_text(getattr(paragraph.paragraph_format, "right_indent", None)),
        "段前": _length_to_text(getattr(paragraph.paragraph_format, "space_before", None)),
        "段后": _length_to_text(getattr(paragraph.paragraph_format, "space_after", None)),
        "行距": _line_spacing_to_text(getattr(paragraph.paragraph_format, "line_spacing", None)),
        "对齐": _alignment_to_text(getattr(paragraph, "alignment", None)),
    }

    mismatches: list[str] = []
    for key in ("首行缩进", "左缩进", "右缩进", "段前", "段后"):
        if not _twips_match(expected[key], actual[key]):
            mismatches.append(key)
    if not _line_spacing_match(expected["行距"], actual["行距"]):
        mismatches.append("行距")
    if expected["对齐"] != actual["对齐"]:
        mismatches.append("对齐")

    if not mismatches:
        return None

    mismatch_label = "、".join(mismatches)
    expected_text = "；".join(
        [
            f"首行缩进={_twips_display(expected['首行缩进'])}",
            f"左缩进={_twips_display(expected['左缩进'])}",
            f"右缩进={_twips_display(expected['右缩进'])}",
            f"段前={_twips_display(expected['段前'])}",
            f"段后={_twips_display(expected['段后'])}",
            f"行距={expected['行距']}",
            f"对齐={expected['对齐']}",
        ]
    )
    actual_text = "；".join(
        [
            f"首行缩进={actual['首行缩进']}",
            f"左缩进={actual['左缩进']}",
            f"右缩进={actual['右缩进']}",
            f"段前={actual['段前']}",
            f"段后={actual['段后']}",
            f"行距={actual['行距']}",
            f"对齐={actual['对齐']}",
        ]
    )
    return {
        "detail": f"{mismatch_label}与模板不一致",
        "expected": expected_text,
        "actual": actual_text,
        "adjustment": f"按模板值调整{mismatch_label}，并清理段落直接格式覆盖。",
    }


def _extract_equation_label(text: str, rule: dict[str, Any]) -> dict[str, Any] | None:
    match = EQUATION_LABEL_REGEX.search(text)
    if not match:
        return None

    chapter = int(match.group("chapter"))
    index = int(match.group("index"))
    separator = match.group("sep")
    left = match.group("left")
    right = match.group("right")
    expected_separator = str(rule.get("separator", "-"))
    expected_left = str(rule.get("left_paren", "("))
    expected_right = str(rule.get("right_paren", ")"))

    return {
        "chapter": chapter,
        "index": index,
        "label": f"{chapter}{separator}{index}",
        "format_ok": separator == expected_separator and left == expected_left and right == expected_right,
        "expected_format": f"{expected_left}{chapter}{expected_separator}{index}{expected_right}",
        "actual_format": f"{left}{chapter}{separator}{index}{right}",
    }


def _extract_figure_label(text: str, rule: dict[str, Any]) -> dict[str, Any] | None:
    match = FIGURE_LABEL_REGEX.match(text)
    if not match:
        return None

    chapter = int(match.group("chapter"))
    index = int(match.group("index"))
    separator = match.group("sep")
    prefix = str(match.group("prefix"))
    expected_separator = str(rule.get("separator", "-"))
    expected_prefix = str(rule.get("prefix", "图"))

    return {
        "chapter": chapter,
        "index": index,
        "label": f"{chapter}{separator}{index}",
        "format_ok": separator == expected_separator and prefix.lower() == expected_prefix.lower(),
        "expected_format": f"{expected_prefix} {chapter}{expected_separator}{index}",
        "actual_format": f"{prefix} {chapter}{separator}{index}",
    }


def _check_equation_numbering(labels: list[dict[str, Any]]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    last_by_chapter: dict[int, int] = {}

    for item in labels:
        paragraph = int(item["paragraph"])
        chapter = int(item["chapter"])
        index = int(item["index"])
        label = str(item["label"])

        if not bool(item["format_ok"]):
            issues.append(
                _issue(
                    category="equation_numbering",
                    paragraph=paragraph,
                    text=label,
                    detail="公式编号格式与模板规则不一致",
                    expected=str(item["expected_format"]),
                    actual=str(item["actual_format"]),
                    adjustment="按模板公式编号格式修改括号和分隔符。",
                )
            )

        last = last_by_chapter.get(chapter, 0)
        expected_next = 1 if last == 0 else last + 1
        if index != expected_next:
            issues.append(
                _issue(
                    category="equation_numbering",
                    paragraph=paragraph,
                    text=label,
                    detail="公式编号顺序不连续",
                    expected=f"{chapter}-{expected_next}",
                    actual=f"{chapter}-{index}",
                    adjustment="按章节内顺序连续重排公式编号。",
                )
            )
        if index > last:
            last_by_chapter[chapter] = index

    return issues


def _check_figure_numbering(labels: list[dict[str, Any]]) -> list[dict[str, Any]]:
    issues: list[dict[str, Any]] = []
    last_by_chapter: dict[int, int] = {}

    for item in labels:
        paragraph = int(item["paragraph"])
        chapter = int(item["chapter"])
        index = int(item["index"])
        label = str(item["label"])

        if not bool(item["format_ok"]):
            issues.append(
                _issue(
                    category="figure_numbering",
                    paragraph=paragraph,
                    text=label,
                    detail="图编号格式与模板规则不一致",
                    expected=str(item["expected_format"]),
                    actual=str(item["actual_format"]),
                    adjustment="按模板图编号格式修改前缀和分隔符。",
                )
            )

        last = last_by_chapter.get(chapter, 0)
        expected_next = 1 if last == 0 else last + 1
        if index != expected_next:
            issues.append(
                _issue(
                    category="figure_numbering",
                    paragraph=paragraph,
                    text=label,
                    detail="图编号顺序不连续",
                    expected=f"{chapter}-{expected_next}",
                    actual=f"{chapter}-{index}",
                    adjustment="按章节内顺序连续重排图编号。",
                )
            )
        if index > last:
            last_by_chapter[chapter] = index

    return issues


def _issue(
    *,
    category: str,
    paragraph: int,
    text: str,
    detail: str,
    expected: str,
    actual: str,
    adjustment: str,
) -> dict[str, Any]:
    severity = CATEGORY_SEVERITY.get(category, "low")
    return {
        "id": f"I{paragraph:04d}-{category}",
        "category": category,
        "category_label": CATEGORY_LABELS.get(category, category),
        "severity": severity,
        "severity_label": SEVERITY_LABELS.get(severity, severity),
        "paragraph": paragraph,
        "location": f"段落#{paragraph}",
        "text": text[:200],
        "detail": detail,
        "expected": expected,
        "actual": actual,
        "adjustment": adjustment,
    }


def _write_compare_report_by_suffix(path: Path, analysis: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    suffix = path.suffix.lower()
    if suffix == ".json":
        path.write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8-sig")
        return
    if suffix in {".txt", ".md"}:
        path.write_text(_render_compare_report_text(analysis), encoding="utf-8-sig")
        return
    if suffix == ".docx":
        _write_compare_report_docx(path, analysis)
        return
    path.write_text(_render_compare_report_text(analysis), encoding="utf-8-sig")


def _render_compare_report_text(analysis: dict[str, Any]) -> str:
    overview = analysis.get("overview", {})
    summary = analysis.get("summary", {})
    categories = analysis.get("categories", [])
    lines: list[str] = []
    lines.append("格式检测报告")
    lines.append(f"最新报告ID: {overview.get('report_id', '')}")
    lines.append(f"检测模板: {overview.get('template_name', '')}")
    lines.append(f"检测文档: {overview.get('source_name', '')}")
    lines.append(f"检测时间: {overview.get('checked_at', '')}")
    lines.append(f"检测状态: {overview.get('status_text', '')}")
    lines.append(f"格式检测错误数: {summary.get('total_issues', 0)}")
    lines.append("")
    lines.append("分类统计:")
    for item in categories:
        lines.append(f"- {item.get('label', '')}: {item.get('count', 0)}")
    lines.append("")
    lines.append("差异明细（含调整建议）:")
    issues = list(analysis.get("issues", []))
    if not issues:
        lines.append("- 未发现差异")
        return "\n".join(lines) + "\n"

    for idx, item in enumerate(issues, start=1):
        lines.append(
            f"{idx}. [{item.get('severity_label', '')}] {item.get('category_label', '')} | "
            f"{item.get('location', '')} | {item.get('detail', '')}"
        )
        lines.append(f"   内容: {item.get('text', '')}")
        lines.append(f"   期望: {item.get('expected', '')}")
        lines.append(f"   实际: {item.get('actual', '')}")
        lines.append(f"   调整建议: {item.get('adjustment', '')}")

    return "\n".join(lines) + "\n"


def _write_compare_report_docx(path: Path, analysis: dict[str, Any]) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    overview = analysis.get("overview", {})
    summary = analysis.get("summary", {})
    categories = analysis.get("categories", [])
    issues = list(analysis.get("issues", []))

    doc = Document()
    doc.add_heading("格式检测报告", level=1)
    doc.add_paragraph(f"最新报告ID：{overview.get('report_id', '')}")
    doc.add_paragraph(f"检测模板：{overview.get('template_name', '')}")
    doc.add_paragraph(f"检测文档：{overview.get('source_name', '')}")
    doc.add_paragraph(f"检测时间：{overview.get('checked_at', '')}")
    doc.add_paragraph(f"检测状态：{overview.get('status_text', '')}")
    doc.add_paragraph(f"格式检测错误数：{summary.get('total_issues', 0)}")

    doc.add_heading("分类统计", level=2)
    stat_table = doc.add_table(rows=1, cols=2)
    stat_header = stat_table.rows[0].cells
    stat_header[0].text = "类别"
    stat_header[1].text = "数量"
    for item in categories:
        row = stat_table.add_row().cells
        row[0].text = str(item.get("label", ""))
        row[1].text = str(item.get("count", 0))

    doc.add_heading("差异明细（含调整建议）", level=2)
    if not issues:
        doc.add_paragraph("未发现差异。")
    else:
        table = doc.add_table(rows=1, cols=8)
        header = table.rows[0].cells
        header[0].text = "编号"
        header[1].text = "级别"
        header[2].text = "类别"
        header[3].text = "位置"
        header[4].text = "问题"
        header[5].text = "期望"
        header[6].text = "实际"
        header[7].text = "调整建议"

        for item in issues:
            row = table.add_row().cells
            row[0].text = str(item.get("id", ""))
            row[1].text = str(item.get("severity_label", ""))
            row[2].text = str(item.get("category_label", ""))
            row[3].text = str(item.get("location", ""))
            row[4].text = str(item.get("detail", ""))
            row[5].text = str(item.get("expected", ""))
            row[6].text = str(item.get("actual", ""))
            row[7].text = str(item.get("adjustment", ""))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))


def _resolve_first_style_name(doc: Any, candidates: tuple[str, ...]) -> str | None:
    names = [str(getattr(style, "name", "") or "") for style in doc.styles]
    direct = {name: name for name in names}
    lowered = {name.lower(): name for name in names}

    for candidate in candidates:
        if candidate in direct:
            return direct[candidate]
    for candidate in candidates:
        matched = lowered.get(candidate.lower())
        if matched:
            return matched
    return None


def _resolve_template_heading_style(doc: Any, level: int) -> str | None:
    style_name = _resolve_first_style_name(
        doc,
        (f"Heading {level}", f"Heading{level}", f"标题 {level}", f"标题{level}"),
    )
    if style_name:
        return style_name

    for style in doc.styles:
        name = str(getattr(style, "name", "") or "")
        lower = name.lower()
        if ("heading" not in lower and "标题" not in name) or not re.search(rf"(?<!\d){level}(?!\d)", name):
            continue
        return name
    return None


def _resolve_template_toc_style(doc: Any, level: int) -> str | None:
    style_name = _resolve_first_style_name(
        doc,
        (f"TOC {level}", f"TOC{level}", f"目录 {level}", f"目录{level}", f"toc {level}", f"toc{level}"),
    )
    if style_name:
        return style_name

    for style in doc.styles:
        name = str(getattr(style, "name", "") or "")
        lower = name.lower()
        if ("toc" not in lower and "目录" not in name) or not re.search(rf"(?<!\d){level}(?!\d)", name):
            continue
        return name
    return None


def _capture_profile(paragraph: Any) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    line_spacing = getattr(fmt, "line_spacing", None)
    line_spacing_kind = "none"
    line_spacing_value: int | float | None = None
    twips = _length_to_twips(line_spacing)
    if twips is not None:
        line_spacing_kind = "exact"
        line_spacing_value = twips
    elif isinstance(line_spacing, (int, float)):
        line_spacing_kind = "multiple"
        line_spacing_value = float(line_spacing)

    return {
        "style": str(getattr(paragraph.style, "name", "") or ""),
        "alignment": _safe_int(getattr(paragraph, "alignment", None)),
        "first_line_indent": _length_to_twips(getattr(fmt, "first_line_indent", None)),
        "left_indent": _length_to_twips(getattr(fmt, "left_indent", None)),
        "right_indent": _length_to_twips(getattr(fmt, "right_indent", None)),
        "space_before": _length_to_twips(getattr(fmt, "space_before", None)),
        "space_after": _length_to_twips(getattr(fmt, "space_after", None)),
        "line_spacing_kind": line_spacing_kind,
        "line_spacing_value": line_spacing_value,
    }


def _heading_level_from_style(paragraph: Any) -> int | None:
    style_name = str(getattr(paragraph.style, "name", "") or "")
    if not style_name:
        return None
    lower = style_name.lower()
    if "heading" not in lower and "标题" not in style_name:
        return None
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 1
    return max(1, min(6, int(match.group(1))))


def _toc_level_from_style(paragraph: Any) -> int | None:
    style_name = str(getattr(paragraph.style, "name", "") or "")
    if not style_name:
        return None
    lower = style_name.lower()
    if "toc" not in lower and "目录" not in style_name:
        return None
    match = re.search(r"(\d+)", style_name)
    if not match:
        return 1
    return max(1, min(6, int(match.group(1))))


def _detect_toc_level(paragraph: Any, text: str) -> int | None:
    from_style = _toc_level_from_style(paragraph)
    if from_style is not None:
        return from_style
    line = re.sub(r"\s+", " ", text.strip())
    if _normalize_heading_text(line) in {"目录", "contents"}:
        return 1
    dot_match = TOC_DOT_REGEX.match(line)
    if not dot_match:
        return None
    title = dot_match.group("title").strip()
    if re.match(r"^\d+[.\uFF0E]\d+[.\uFF0E]\d+", title):
        return 3
    if re.match(r"^\d+[.\uFF0E]\d+", title):
        return 2
    return 1


def _detect_heading_level(paragraph: Any, text: str) -> int | None:
    style_level = _heading_level_from_style(paragraph)
    if style_level is not None:
        return style_level

    line = re.sub(r"\s+", " ", text.strip())
    if not line or len(line) > 80:
        return None
    if line.startswith("[") and re.match(r"^\[\d+\]", line):
        return None
    if re.search(r"[。！？!?；;]$", line):
        return None
    if re.match(r"^\s*第[\u96f6\u3007\u4e00-\u9fa5\d]+\s*章(?:[\s:：.\-]+.+)?$", line):
        return 1
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+\s*[.\uFF0E]\s*\d+\s+.+$", line):
        return 3
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+\s+.+$", line):
        return 2
    if re.match(r"^\s*\d{1,2}(?:\s*[.\uFF0E、)]\s*|\s+).+$", line):
        return 1
    return None


def _is_reference_heading_text(text: str) -> bool:
    return _normalize_heading_text(text) in REFERENCE_HEADINGS


def _is_reference_entry_text(paragraph: Any, text: str) -> bool:
    line = text.strip()
    if re.match(r"^\[\d+\]", line):
        return True
    if re.match(r"^\d+[.)、]\s*", line):
        return True
    first_line = _length_to_twips(getattr(paragraph.paragraph_format, "first_line_indent", None))
    return first_line is not None and first_line < 0


def _is_toc_intro_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return any(keyword in compact for keyword in ("目录", "章节安排", "具体安排如下", "拟分为"))


def _infer_manual_toc_level_from_text(text: str) -> int | None:
    line = re.sub(r"\s+", " ", text.strip())
    if not line:
        return None
    if re.match(r"^\s*第[\u96f6\u3007\u4e00-\u9fa5\d]+\s*章", line):
        return 1
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+\s*[.\uFF0E]\s*\d+\s+.+$", line):
        return 3
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+\s+.+$", line):
        return 2
    return None


def _is_body_text_candidate(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    if len(compact) < 12:
        return False
    if _infer_manual_toc_level_from_text(text) is not None:
        return False
    if _is_reference_heading_text(text):
        return False
    return True


def _is_better_reference_profile(candidate: dict[str, Any], existing: dict[str, Any]) -> bool:
    candidate_first = _safe_int(candidate.get("first_line_indent"))
    existing_first = _safe_int(existing.get("first_line_indent"))
    if candidate_first is not None and candidate_first < 0:
        if existing_first is None or existing_first >= 0:
            return True
    return False


def _length_to_twips(length_value: Any) -> int | None:
    if length_value is None:
        return None
    twips = getattr(length_value, "twips", None)
    if twips is None:
        return None
    try:
        return int(twips)
    except Exception:  # noqa: BLE001
        return None


def _length_to_text(length_value: Any) -> str:
    twips = _length_to_twips(length_value)
    if twips is None:
        return "None"
    return f"{twips} twips"


def _line_spacing_to_text(value: Any) -> str:
    if value is None:
        return "None"
    if isinstance(value, (int, float)):
        return f"{float(value):.2f}"
    twips = _length_to_twips(value)
    if twips is None:
        return str(value)
    return f"{twips} twips"


def _expected_line_spacing_text(profile: dict[str, Any]) -> str:
    kind = str(profile.get("line_spacing_kind", "none"))
    value = profile.get("line_spacing_value")
    if kind == "multiple" and isinstance(value, (int, float)):
        return f"{float(value):.2f}"
    if kind == "exact":
        number = _safe_int(value)
        return f"{number} twips" if number is not None else "None"
    return "None"


def _alignment_to_text(value: Any) -> str:
    numeric = _safe_int(value)
    if numeric is None:
        return "None"
    mapping = {0: "left", 1: "center", 2: "right", 3: "justify"}
    return mapping.get(numeric, str(numeric))


def _twips_match(expected_twips: Any, actual_text: str) -> bool:
    expected = _safe_int(expected_twips)
    actual = _parse_twips_text(actual_text)
    if expected is None and actual is None:
        return True
    if expected is None or actual is None:
        return False
    return abs(expected - actual) <= 20


def _line_spacing_match(expected_text: str, actual_text: str) -> bool:
    if expected_text == "None" and actual_text == "None":
        return True
    expected_multiple = _parse_multiple(expected_text)
    actual_multiple = _parse_multiple(actual_text)
    if expected_multiple is not None and actual_multiple is not None:
        return abs(expected_multiple - actual_multiple) <= 0.05

    expected_twips = _parse_twips_text(expected_text)
    actual_twips = _parse_twips_text(actual_text)
    if expected_twips is None and actual_twips is None:
        return expected_text == actual_text
    if expected_twips is None or actual_twips is None:
        return False
    return abs(expected_twips - actual_twips) <= 20


def _twips_display(value: Any) -> str:
    number = _safe_int(value)
    if number is None:
        return "None"
    return f"{number} twips"


def _parse_twips_text(text: str) -> int | None:
    match = re.search(r"(-?\d+)\s*twips", text)
    if not match:
        return None
    return int(match.group(1))


def _parse_multiple(text: str) -> float | None:
    if text.endswith("twips"):
        return None
    try:
        return float(text)
    except Exception:  # noqa: BLE001
        return None


def _safe_int(value: Any) -> int | None:
    if value is None:
        return None
    try:
        return int(value)
    except Exception:  # noqa: BLE001
        return None


def _normalize_heading_text(text: str) -> str:
    return re.sub(r"\s+", "", text.strip().lower())


def _build_result(
    operation: str,
    source: Path,
    destination: Path | None,
    status: OperationStatus,
    message: str,
    started: datetime,
    started_at: str,
) -> OperationResult:
    finished = datetime.now()
    finished_at = now_iso()
    return OperationResult(
        operation=operation,
        source=str(source),
        destination=str(destination) if destination else None,
        status=status,
        message=message,
        started_at=started_at,
        finished_at=finished_at,
        duration_ms=duration_ms(started, finished),
    )
