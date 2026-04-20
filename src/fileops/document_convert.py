from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from .models import OperationResult, OperationStatus
from .utils import duration_ms, ensure_workspace_path, now_iso, unique_path

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover
    PdfReader = None


SUPPORTED_FORMATS = {"docx", "pdf", "markdown"}
MARKDOWN_EXTENSIONS = {".md", ".markdown"}
FOLIO_CLI_ENV = "FILEOPS_FOLIO_CLI"
FOLIO_PROJECT_ENV = "FILEOPS_FOLIO_PROJECT"
FOLIO_STRICT_ENV = "FILEOPS_FOLIO_STRICT"


def convert_documents_format(
    sources: list[Path],
    destination: Path,
    workspace: Path,
    dry_run: bool,
    source_format: str,
    target_format: str,
) -> list[OperationResult]:
    if source_format not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported source format: {source_format}")
    if target_format not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported target format: {target_format}")
    if source_format == target_format:
        raise ValueError("Source format and target format must be different.")

    destination = destination.resolve(strict=False)
    ensure_workspace_path(destination, workspace)
    if not dry_run:
        destination.mkdir(parents=True, exist_ok=True)

    results: list[OperationResult] = []
    for source in sources:
        started = datetime.now()
        started_at = now_iso()
        conversion_engine: str | None = None

        try:
            ensure_workspace_path(source, workspace)
            if not source.exists():
                raise FileNotFoundError(f"Source does not exist: {source}")
            if source.is_dir():
                raise IsADirectoryError(f"Document convert supports files only: {source}")
            if not _matches_source_format(source, source_format):
                raise ValueError(f"Source format does not match convert source format setting: {source.name}")

            output_path = destination / f"{source.stem}_converted.{target_format}"
            if output_path.exists():
                output_path = unique_path(output_path)

            if dry_run:
                message = f"Would convert {source_format.upper()} to {target_format.upper()} -> {output_path.name}"
                results.append(
                    _build_result("doc_convert", source, output_path, OperationStatus.DRY_RUN, message, started, started_at)
                )
                continue

            if source_format == "pdf" and target_format == "docx":
                _convert_pdf_to_docx(source, output_path)
            elif source_format == "docx" and target_format == "pdf":
                _convert_docx_to_pdf(source, output_path)
            elif source_format == "markdown" and target_format == "pdf":
                conversion_engine = _convert_markdown_to_pdf(source, output_path, workspace)
            elif source_format == "markdown" and target_format == "docx":
                conversion_engine = _convert_markdown_to_docx(source, output_path, workspace)
            else:
                raise ValueError(f"Unsupported conversion pair: {source_format} -> {target_format}")

            message = f"Converted {source_format.upper()} to {target_format.upper()} -> {output_path.name}"
            if conversion_engine:
                message = f"{message} (engine: {conversion_engine})"
            results.append(_build_result("doc_convert", source, output_path, OperationStatus.SUCCESS, message, started, started_at))
        except Exception as exc:  # noqa: BLE001
            results.append(_build_result("doc_convert", source, None, OperationStatus.FAILED, str(exc), started, started_at))

    return results


def _matches_source_format(source: Path, source_format: str) -> bool:
    ext = source.suffix.lower()
    if source_format == "markdown":
        return ext in MARKDOWN_EXTENSIONS
    return ext == f".{source_format}"


def _convert_pdf_to_docx(source: Path, output: Path) -> None:
    if PdfReader is None or Document is None:
        raise RuntimeError("pypdf/python-docx is not installed. Please install dependencies first.")

    reader = PdfReader(str(source))
    doc = Document()

    has_text = False
    for page_index, page in enumerate(reader.pages, start=1):
        raw_text = page.extract_text() or ""
        lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
        if not lines:
            continue
        if has_text and page_index > 1:
            doc.add_page_break()
        for line in lines:
            doc.add_paragraph(line)
        has_text = True

    if not has_text:
        doc.add_paragraph("No extractable text found in source PDF.")

    doc.save(str(output))


def _convert_docx_to_pdf(source: Path, output: Path) -> None:
    ps_source = _ps_quote(str(source))
    ps_output = _ps_quote(str(output))
    script = (
        "$ErrorActionPreference='Stop';"
        f"$src={ps_source};"
        f"$dst={ps_output};"
        "$word=$null;$doc=$null;"
        "try{"
        "$word=New-Object -ComObject Word.Application;"
        "$word.Visible=$false;"
        "$word.DisplayAlerts=0;"
        "$doc=$word.Documents.Open($src,$false,$true);"
        "$doc.ExportAsFixedFormat($dst,17);"
        "}finally{"
        "if($doc -ne $null){$doc.Close($false);[void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($doc)};"
        "if($word -ne $null){$word.Quit();[void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($word)};"
        "}"
    )

    completed = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-ExecutionPolicy", "Bypass", "-Command", script],
        capture_output=True,
        text=True,
        check=False,
    )
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout or "").strip()
        raise RuntimeError(
            "Microsoft Word COM export failed. Please ensure Microsoft Word is installed and can open this DOCX."
            + (f" Detail: {detail}" if detail else "")
        )
    if not output.exists():
        raise RuntimeError("Word export finished but target PDF was not created.")


def _convert_markdown_to_pdf(source: Path, output: Path, workspace: Path) -> str:
    with TemporaryDirectory() as temp_dir:
        temp_docx = Path(temp_dir) / f"{source.stem}_markdown_bridge.docx"
        conversion_engine = _convert_markdown_to_docx(source, temp_docx, workspace)
        _convert_docx_to_pdf(temp_docx, output)
    return conversion_engine


def _convert_markdown_to_docx(source: Path, output: Path, workspace: Path) -> str:
    if _try_convert_markdown_with_folio(source, output, workspace):
        return "folio"

    _convert_markdown_to_docx_builtin(source, output)
    return "builtin"


def _convert_markdown_to_docx_builtin(source: Path, output: Path) -> None:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    lines = source.read_text(encoding="utf-8", errors="ignore").splitlines()
    doc = Document()
    _append_markdown_lines_to_doc(doc, lines)
    if not doc.paragraphs and not doc.tables:
        doc.add_paragraph("")
    doc.save(str(output))


def _try_convert_markdown_with_folio(source: Path, output: Path, workspace: Path) -> bool:
    commands = _build_folio_command_candidates(source, output, workspace)
    strict_mode = os.getenv(FOLIO_STRICT_ENV, "").strip().lower() in {"1", "true", "yes", "on"}
    errors: list[str] = []

    for command, cwd in commands:
        try:
            completed = subprocess.run(
                command,
                capture_output=True,
                text=True,
                check=False,
                cwd=str(cwd) if cwd is not None else None,
            )
        except OSError as exc:
            errors.append(f"{' '.join(command)} -> {exc}")
            continue

        if completed.returncode == 0 and output.exists():
            return True

        detail = (completed.stderr or completed.stdout or "").strip()
        if detail:
            errors.append(f"{' '.join(command)} -> {detail}")
        else:
            errors.append(f"{' '.join(command)} -> exit code {completed.returncode}")

    if strict_mode and errors:
        error_text = "; ".join(errors[:2])
        raise RuntimeError(f"Folio markdown conversion failed. {error_text}")

    return False


def _build_folio_command_candidates(source: Path, output: Path, workspace: Path) -> list[tuple[list[str], Path | None]]:
    source_text = str(source)
    output_text = str(output)
    command_candidates: list[tuple[list[str], Path | None]] = []
    seen: set[tuple[tuple[str, ...], str | None]] = set()
    has_cargo = shutil.which("cargo") is not None

    def _append(command: list[str], cwd: Path | None) -> None:
        marker = (tuple(command), str(cwd) if cwd is not None else None)
        if marker in seen:
            return
        seen.add(marker)
        command_candidates.append((command, cwd))

    configured_cli = os.getenv(FOLIO_CLI_ENV, "").strip()
    if configured_cli:
        _append([configured_cli, source_text, "-o", output_text], None)

    for bundled_executable in _candidate_bundled_folio_executables(workspace, source):
        if bundled_executable.exists():
            _append([str(bundled_executable), source_text, "-o", output_text], None)

    cli_on_path = shutil.which("scribe-cli")
    if cli_on_path:
        _append([cli_on_path, source_text, "-o", output_text], None)

    for project_root in _candidate_folio_project_roots(workspace, source):
        for executable in _candidate_folio_executables(project_root):
            if executable.exists():
                _append([str(executable), source_text, "-o", output_text], None)

        if has_cargo:
            _append(
                ["cargo", "run", "-p", "scribe-cli", "--", source_text, "-o", output_text],
                project_root,
            )

    return command_candidates


def _candidate_folio_project_roots(workspace: Path, source: Path) -> list[Path]:
    roots: list[Path] = []
    seen: set[Path] = set()

    def _append(path: Path) -> None:
        resolved = path.resolve(strict=False)
        if resolved in seen:
            return
        cargo_file = resolved / "Cargo.toml"
        if not cargo_file.exists():
            return
        if not (resolved / "scribe-cli").exists():
            return
        seen.add(resolved)
        roots.append(resolved)

    configured_project = os.getenv(FOLIO_PROJECT_ENV, "").strip()
    if configured_project:
        _append(Path(configured_project))

    _append(workspace / "Folio-master")
    _append(source.parent / "Folio-master")
    _append(Path(__file__).resolve().parents[2] / "Folio-master")
    _append(Path.cwd() / "Folio-master")

    return roots


def _candidate_folio_executables(project_root: Path) -> list[Path]:
    executable_name = _folio_executable_name()
    return [
        project_root / "target" / "release" / executable_name,
        project_root / "target" / "debug" / executable_name,
    ]


def _candidate_bundled_folio_executables(workspace: Path, source: Path) -> list[Path]:
    executable_name = _folio_executable_name()
    candidates: list[Path] = []
    seen: set[Path] = set()

    def _append(path: Path) -> None:
        resolved = path.resolve(strict=False)
        if resolved in seen:
            return
        seen.add(resolved)
        candidates.append(resolved)

    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        _append(Path(str(meipass)) / "folio" / executable_name)

    executable_path = Path(sys.executable).resolve(strict=False)
    _append(executable_path.parent / "folio" / executable_name)

    repo_root = Path(__file__).resolve().parents[2]
    _append(repo_root / "vendor" / "folio" / "bin" / executable_name)
    _append(Path.cwd() / "vendor" / "folio" / "bin" / executable_name)
    _append(workspace / "vendor" / "folio" / "bin" / executable_name)
    _append(source.parent / "vendor" / "folio" / "bin" / executable_name)

    return candidates


def _folio_executable_name() -> str:
    return "scribe-cli.exe" if os.name == "nt" else "scribe-cli"


def _append_markdown_lines_to_doc(doc: Any, lines: list[str]) -> None:
    heading_regex = re.compile(r"^(#{1,6})\s+(.*)$")
    bullet_regex = re.compile(r"^[-*+]\s+(.+)$")
    ordered_regex = re.compile(r"^\d+\.\s+(.+)$")

    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if _looks_like_markdown_table_row(line):
            table_lines = [line]
            index += 1
            while index < len(lines) and _looks_like_markdown_table_row(lines[index].rstrip()):
                table_lines.append(lines[index].rstrip())
                index += 1
            _append_markdown_table_to_doc(doc, table_lines)
            continue

        heading_match = heading_regex.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            doc.add_heading(text if text else line, level=level)
            index += 1
            continue

        bullet_match = bullet_regex.match(line)
        if bullet_match:
            _add_paragraph_with_optional_style(doc, bullet_match.group(1).strip(), "List Bullet")
            index += 1
            continue

        ordered_match = ordered_regex.match(line)
        if ordered_match:
            _add_paragraph_with_optional_style(doc, ordered_match.group(1).strip(), "List Number")
            index += 1
            continue

        doc.add_paragraph(line)
        index += 1


def _add_paragraph_with_optional_style(doc: Any, text: str, style_name: str) -> None:
    try:
        doc.add_paragraph(text, style=style_name)
    except Exception:  # noqa: BLE001
        doc.add_paragraph(text)


def _looks_like_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _append_markdown_table_to_doc(doc: Any, table_lines: list[str]) -> None:
    rows = [_parse_markdown_table_row(item) for item in table_lines]
    if not rows:
        return

    if len(rows) >= 2 and _is_markdown_separator_row(rows[1]):
        rows = [rows[0], *rows[2:]]
    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)

    for row_idx, row in enumerate(rows):
        normalized = row + [""] * (cols - len(row))
        for col_idx, value in enumerate(normalized):
            table.cell(row_idx, col_idx).text = value


def _parse_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _is_markdown_separator_row(row: list[str]) -> bool:
    for cell in row:
        normalized = cell.replace(" ", "")
        if not normalized:
            continue
        if not re.fullmatch(r":?-{3,}:?", normalized):
            return False
    return True


def _ps_quote(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def _build_result(
    operation: str,
    source: Path,
    destination: Path | None,
    status: OperationStatus,
    message: str,
    started: datetime,
    started_at: str,
) -> OperationResult:
    finished = datetime.now()
    return OperationResult(
        operation=operation,
        source=str(source),
        destination=str(destination) if destination is not None else None,
        status=status,
        message=message,
        started_at=started_at,
        finished_at=now_iso(),
        duration_ms=duration_ms(started, finished),
    )
