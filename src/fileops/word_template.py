from __future__ import annotations

import os
import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path
from typing import Any

from .models import OperationResult, OperationStatus
from .utils import duration_ms, ensure_workspace_path, now_iso, unique_path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt, Twips
except ImportError:  # pragma: no cover
    Document = None
    qn = None
    Pt = None
    Twips = None


def template_library_dir() -> Path:
    candidates: list[Path] = []

    local_app_data = os.getenv("LOCALAPPDATA")
    if local_app_data:
        candidates.append(Path(local_app_data) / "FileOps" / "templates")

    roaming_app_data = os.getenv("APPDATA")
    if roaming_app_data:
        candidates.append(Path(roaming_app_data) / "FileOps" / "templates")

    candidates.append(Path.home() / ".fileops" / "templates")
    candidates.append(Path.cwd() / ".fileops" / "templates")

    for candidate in candidates:
        try:
            candidate.mkdir(parents=True, exist_ok=True)
            return candidate
        except OSError:
            continue

    raise PermissionError("Unable to create template library directory. Check folder permissions.")


def list_word_templates() -> list[Path]:
    base = template_library_dir()
    return sorted([item for item in base.glob("*.docx") if item.is_file()], key=lambda item: item.name.lower())


def import_word_template(template_file: Path) -> Path:
    template_file = template_file.resolve(strict=False)
    if not template_file.exists():
        raise FileNotFoundError(f"Template does not exist: {template_file}")
    if template_file.suffix.lower() != ".docx":
        raise ValueError("Only .docx files can be imported as templates.")

    target_dir = template_library_dir()
    target = target_dir / template_file.name
    if target.exists():
        target = unique_path(target)
    target.write_bytes(template_file.read_bytes())
    return target


def format_word_documents(
    sources: list[Path],
    destination: Path,
    workspace: Path,
    dry_run: bool,
    template_path: Path,
) -> list[OperationResult]:
    if Document is None or qn is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    destination = destination.resolve(strict=False)
    ensure_workspace_path(destination, workspace)
    if not dry_run:
        destination.mkdir(parents=True, exist_ok=True)

    template_path = template_path.resolve(strict=False)
    if not template_path.exists():
        raise FileNotFoundError(f"Template does not exist: {template_path}")
    if template_path.suffix.lower() != ".docx":
        raise ValueError("Template must be a .docx file.")

    results: list[OperationResult] = []
    for source in sources:
        started = datetime.now()
        started_at = now_iso()
        try:
            ensure_workspace_path(source, workspace)
            if not source.exists():
                raise FileNotFoundError(f"Source does not exist: {source}")
            if source.is_dir():
                raise IsADirectoryError(f"Word format supports files only: {source}")
            if source.suffix.lower() != ".docx":
                raise ValueError("Word format supports .docx files only.")

            output_path = destination / f"{source.stem}_formatted.docx"
            if output_path.exists():
                output_path = unique_path(output_path)

            if dry_run:
                message = f"Would format by template '{template_path.name}' -> {output_path.name}"
                results.append(_build_result("word_format", source, output_path, OperationStatus.DRY_RUN, message, started, started_at))
                continue

            _apply_template_format(source=source, template=template_path, output=output_path)
            message = f"Formatted by template '{template_path.name}' -> {output_path.name}"
            results.append(_build_result("word_format", source, output_path, OperationStatus.SUCCESS, message, started, started_at))
        except Exception as exc:  # noqa: BLE001
            results.append(_build_result("word_format", source, None, OperationStatus.FAILED, str(exc), started, started_at))

    return results


def apply_template_format_to_document(source: Path, template: Path, output: Path) -> None:
    _apply_template_format(source=source, template=template, output=output)


def _apply_template_format(source: Path, template: Path, output: Path) -> None:
    template_doc = Document(str(template))
    source_doc = Document(str(source))

    style_map = _build_template_style_map(template_doc)
    profile_map = _build_template_paragraph_profile_map(template_doc)
    run_profile_map = _build_template_run_profile_map(template_doc)
    _clear_template_body_keep_sectpr(template_doc)

    table_style = _resolve_first_table_style(template_doc)
    context: dict[str, Any] = {"toc_mode": False, "in_references": False}

    source_blocks = _iter_body_blocks(source_doc)
    for block in source_blocks:
        if block.tag == qn("w:p"):
            paragraph = _paragraph_from_xml(block, source_doc)
            _copy_paragraph_to_template(
                source_paragraph=paragraph,
                target_doc=template_doc,
                style_map=style_map,
                profile_map=profile_map,
                run_profile_map=run_profile_map,
                context=context,
            )
            continue

        if block.tag == qn("w:tbl"):
            table = _table_from_xml(block, source_doc)
            _copy_table_to_template(
                source_table=table,
                target_doc=template_doc,
                table_style=table_style,
                paragraph_profile=profile_map.get("normal"),
                run_profile=run_profile_map.get("normal"),
            )

    template_doc.save(str(output))


def _clear_template_body_keep_sectpr(doc: Any) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    sectpr = next((child for child in children if child.tag == qn("w:sectPr")), None)

    for child in children:
        body.remove(child)

    if sectpr is not None:
        body.append(sectpr)
    else:
        doc.add_paragraph("")


def _resolve_first_table_style(doc: Any) -> str | None:
    for style in doc.styles:
        style_type = str(getattr(style.type, "name", "") or "")
        if style_type.lower() == "table":
            return style.name
    return None


def _build_template_style_map(doc: Any) -> dict[str, Any]:
    normal_style = _resolve_first_available_style_name(doc, ("Normal", "正文", "Body Text"))
    if normal_style is None:
        normal_style = "Normal"

    heading_styles: dict[int, str] = {}
    toc_styles: dict[int, str] = {}
    for level in range(1, 7):
        heading_name = _resolve_heading_style_in_template(doc, level)
        if heading_name:
            heading_styles[level] = heading_name
        toc_name = _resolve_toc_style_in_template(doc, level)
        if toc_name:
            toc_styles[level] = toc_name

    for paragraph in doc.paragraphs:
        style_name = str(getattr(paragraph.style, "name", "") or "")
        if not style_name:
            continue
        text = str(paragraph.text or "").strip()
        if not text:
            continue

        guessed_toc = _infer_toc_level_from_text(text)
        if guessed_toc is not None and guessed_toc not in toc_styles:
            toc_styles[guessed_toc] = style_name

        guessed_heading = _infer_heading_level_from_text(text)
        if guessed_heading is not None and guessed_heading not in heading_styles:
            heading_styles[guessed_heading] = style_name

    reference_style = normal_style
    reference_heading_style = heading_styles.get(1, normal_style)
    in_reference_section = False
    for paragraph in doc.paragraphs:
        text = str(paragraph.text or "").strip()
        if not text:
            continue
        if _is_reference_heading_text(text):
            in_reference_section = True
            heading_style_name = str(getattr(paragraph.style, "name", "") or "")
            if heading_style_name:
                reference_heading_style = heading_style_name
            continue
        if in_reference_section and _is_reference_entry_text(paragraph, text):
            reference_style_name = str(getattr(paragraph.style, "name", "") or "")
            if reference_style_name:
                reference_style = reference_style_name
            break

    return {
        "normal": normal_style,
        "heading": heading_styles,
        "toc": toc_styles,
        "reference": reference_style,
        "reference_heading": reference_heading_style,
    }


def _resolve_first_available_style_name(doc: Any, candidates: tuple[str, ...]) -> str | None:
    names = {str(style.name): style for style in doc.styles}
    for candidate in candidates:
        if candidate in names:
            return candidate

    lowered = {name.lower(): name for name in names}
    for candidate in candidates:
        matched = lowered.get(candidate.lower())
        if matched:
            return matched
    return None


def _resolve_heading_style_in_template(doc: Any, level: int) -> str | None:
    candidates = (
        f"Heading {level}",
        f"Heading{level}",
        f"标题 {level}",
        f"标题{level}",
    )
    style_name = _resolve_first_available_style_name(doc, candidates)
    if style_name:
        return style_name

    for style in doc.styles:
        name = str(getattr(style, "name", "") or "")
        lower = name.lower()
        if ("heading" not in lower and "标题" not in name) or not re.search(rf"(?<!\d){level}(?!\d)", name):
            continue
        return name
    return None


def _resolve_toc_style_in_template(doc: Any, level: int) -> str | None:
    candidates = (
        f"TOC {level}",
        f"TOC{level}",
        f"目录 {level}",
        f"目录{level}",
    )
    style_name = _resolve_first_available_style_name(doc, candidates)
    if style_name:
        return style_name

    for style in doc.styles:
        name = str(getattr(style, "name", "") or "")
        lower = name.lower()
        if ("toc" not in lower and "目录" not in name) or not re.search(rf"(?<!\d){level}(?!\d)", name):
            continue
        return name
    return None


def _build_template_paragraph_profile_map(doc: Any) -> dict[str, Any]:
    profiles: dict[str, Any] = {
        "normal": None,
        "heading": {},
        "toc": {},
        "reference_heading": None,
        "reference_entry": None,
    }
    context: dict[str, bool] = {"toc_mode": False, "in_references": False}
    seen_heading_style = False

    for paragraph in doc.paragraphs:
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text:
            continue

        if _is_reference_heading_text(text):
            context["in_references"] = True
            context["toc_mode"] = False
            if profiles["reference_heading"] is None:
                profiles["reference_heading"] = _capture_paragraph_profile(paragraph)
            continue

        if _is_toc_intro_text(text):
            context["toc_mode"] = True

        if bool(context.get("toc_mode")):
            toc_level = _infer_manual_toc_level_from_text(text) or _infer_toc_level_from_text(text)
            if toc_level is not None and toc_level not in profiles["toc"]:
                profiles["toc"][toc_level] = _capture_paragraph_profile(paragraph)
                continue
            if toc_level is None:
                context["toc_mode"] = False

        if bool(context.get("in_references")) and _is_reference_entry_text(paragraph, text):
            existing = profiles["reference_entry"]
            candidate = _capture_paragraph_profile(paragraph)
            if existing is None or _is_better_reference_profile(candidate, existing):
                profiles["reference_entry"] = candidate
            continue

        heading_level = _resolve_heading_level_from_style(paragraph)
        if heading_level is not None and heading_level not in profiles["heading"]:
            profiles["heading"][heading_level] = _capture_paragraph_profile(paragraph)
            seen_heading_style = True
            continue
        if heading_level is not None:
            seen_heading_style = True

        if profiles["normal"] is None and seen_heading_style and _is_body_text_candidate(text):
            profiles["normal"] = _capture_paragraph_profile(paragraph)

    if profiles["normal"] is None:
        for paragraph in doc.paragraphs:
            text = str(getattr(paragraph, "text", "") or "").strip()
            if not text:
                continue
            if _is_body_text_candidate(text):
                profiles["normal"] = _capture_paragraph_profile(paragraph)
                break

    if profiles["normal"] is None:
        for paragraph in doc.paragraphs:
            text = str(getattr(paragraph, "text", "") or "").strip()
            if not text:
                continue
            profiles["normal"] = _capture_paragraph_profile(paragraph)
            break

    if profiles["reference_heading"] is None:
        profiles["reference_heading"] = profiles["heading"].get(1) or profiles["normal"]
    if profiles["reference_entry"] is None:
        profiles["reference_entry"] = profiles["normal"]
    return profiles


def _build_template_run_profile_map(doc: Any) -> dict[str, Any]:
    run_profiles: dict[str, Any] = {
        "normal": None,
        "heading": {},
        "toc": {},
        "reference_heading": None,
        "reference_entry": None,
    }
    context: dict[str, bool] = {"toc_mode": False, "in_references": False}

    for paragraph in doc.paragraphs:
        text = str(getattr(paragraph, "text", "") or "").strip()
        if not text:
            continue
        sample_run = _find_first_text_run(paragraph)
        if sample_run is None:
            continue
        captured = _capture_run_profile(sample_run)

        role, level = _classify_source_paragraph(paragraph, context)
        if role == "heading" and level is not None and level not in run_profiles["heading"]:
            run_profiles["heading"][level] = captured
        elif role == "toc" and level is not None and level not in run_profiles["toc"]:
            run_profiles["toc"][level] = captured
        elif role == "reference_heading" and run_profiles["reference_heading"] is None:
            run_profiles["reference_heading"] = captured
        elif role == "reference_entry" and run_profiles["reference_entry"] is None:
            run_profiles["reference_entry"] = captured
        elif role == "normal" and run_profiles["normal"] is None:
            run_profiles["normal"] = captured

    if run_profiles["normal"] is None:
        run_profiles["normal"] = {}
    if run_profiles["reference_heading"] is None:
        run_profiles["reference_heading"] = run_profiles["heading"].get(1) or run_profiles["normal"]
    if run_profiles["reference_entry"] is None:
        run_profiles["reference_entry"] = run_profiles["normal"]
    return run_profiles


def _find_first_text_run(paragraph: Any) -> Any | None:
    for run in paragraph.runs:
        if str(getattr(run, "text", "") or "").strip():
            return run
    return None


def _capture_run_profile(run: Any) -> dict[str, Any]:
    font = run.font
    east_asia = None
    ascii_name = None
    hansi_name = None
    if qn is not None:
        rfonts = run._element.rPr.rFonts if run._element.rPr is not None else None
        if rfonts is not None:
            east_asia = rfonts.get(qn("w:eastAsia"))
            ascii_name = rfonts.get(qn("w:ascii"))
            hansi_name = rfonts.get(qn("w:hAnsi"))

    size_pt = None
    if font.size is not None:
        try:
            size_pt = float(font.size.pt)
        except Exception:  # noqa: BLE001
            size_pt = None

    color_rgb = None
    if font.color is not None and getattr(font.color, "rgb", None) is not None:
        color_rgb = str(font.color.rgb)

    return {
        "name": font.name,
        "east_asia": east_asia,
        "ascii": ascii_name,
        "hansi": hansi_name,
        "size_pt": size_pt,
        "bold": font.bold,
        "italic": font.italic,
        "underline": font.underline,
        "color_rgb": color_rgb,
        "highlight_color": font.highlight_color,
    }


def _resolve_profile_for_role(profile_map: dict[str, Any], role: str, level: int | None) -> dict[str, Any] | None:
    if role == "reference_entry":
        return profile_map.get("reference_entry")
    if role == "reference_heading":
        return profile_map.get("reference_heading")
    if role == "toc":
        toc_map = profile_map.get("toc", {})
        if level is not None and level in toc_map:
            return toc_map[level]
        return toc_map.get(1)
    if role == "heading":
        heading_map = profile_map.get("heading", {})
        if level is not None and level in heading_map:
            return heading_map[level]
        return heading_map.get(1)
    return profile_map.get("normal")


def _resolve_run_profile_for_role(run_profile_map: dict[str, Any], role: str, level: int | None) -> dict[str, Any] | None:
    if role == "reference_entry":
        return run_profile_map.get("reference_entry")
    if role == "reference_heading":
        return run_profile_map.get("reference_heading")
    if role == "toc":
        toc_map = run_profile_map.get("toc", {})
        if level is not None and level in toc_map:
            return toc_map[level]
        return toc_map.get(1) or run_profile_map.get("normal")
    if role == "heading":
        heading_map = run_profile_map.get("heading", {})
        if level is not None and level in heading_map:
            return heading_map[level]
        return heading_map.get(1) or run_profile_map.get("normal")
    return run_profile_map.get("normal")


def _capture_paragraph_profile(paragraph: Any) -> dict[str, Any]:
    fmt = paragraph.paragraph_format
    line_spacing = getattr(fmt, "line_spacing", None)
    line_spacing_kind = "none"
    line_spacing_value: int | float | None = None
    twips = _length_to_twips(line_spacing)
    if twips is not None:
        line_spacing_kind = "exact"
        line_spacing_value = twips
    elif isinstance(line_spacing, (int, float)):
        line_spacing_kind = "multiple"
        line_spacing_value = float(line_spacing)

    return {
        "style": str(getattr(paragraph.style, "name", "") or ""),
        "alignment": getattr(paragraph, "alignment", None),
        "left_indent": _length_to_twips(getattr(fmt, "left_indent", None)),
        "right_indent": _length_to_twips(getattr(fmt, "right_indent", None)),
        "first_line_indent": _length_to_twips(getattr(fmt, "first_line_indent", None)),
        "space_before": _length_to_twips(getattr(fmt, "space_before", None)),
        "space_after": _length_to_twips(getattr(fmt, "space_after", None)),
        "line_spacing_kind": line_spacing_kind,
        "line_spacing_value": line_spacing_value,
        "line_spacing_rule": getattr(fmt, "line_spacing_rule", None),
    }


def _apply_profile_format(paragraph: Any, profile: dict[str, Any] | None) -> None:
    if not profile:
        return
    style_name = str(profile.get("style", "") or "")
    if style_name:
        try:
            paragraph.style = style_name
        except Exception:  # noqa: BLE001
            pass

    paragraph.alignment = profile.get("alignment")
    fmt = paragraph.paragraph_format
    fmt.left_indent = _twips_to_length(profile.get("left_indent"))
    fmt.right_indent = _twips_to_length(profile.get("right_indent"))
    fmt.first_line_indent = _twips_to_length(profile.get("first_line_indent"))
    fmt.space_before = _twips_to_length(profile.get("space_before"))
    fmt.space_after = _twips_to_length(profile.get("space_after"))

    line_spacing_kind = profile.get("line_spacing_kind")
    line_spacing_value = profile.get("line_spacing_value")
    if line_spacing_kind == "multiple" and isinstance(line_spacing_value, (int, float)):
        fmt.line_spacing = float(line_spacing_value)
    elif line_spacing_kind == "exact":
        fmt.line_spacing = _twips_to_length(line_spacing_value)
    else:
        fmt.line_spacing = None
    fmt.line_spacing_rule = profile.get("line_spacing_rule")


def _length_to_twips(length_value: Any) -> int | None:
    if length_value is None:
        return None
    twips = getattr(length_value, "twips", None)
    if twips is None:
        return None
    try:
        return int(twips)
    except Exception:  # noqa: BLE001
        return None


def _twips_to_length(value: Any) -> Any:
    if value is None:
        return None
    try:
        if Twips is not None:
            return Twips(int(value))
    except Exception:  # noqa: BLE001
        return None
    return None


def _resolve_heading_level_from_style(paragraph: Any) -> int | None:
    try:
        style_name = str(getattr(paragraph.style, "name", "") or "")
    except Exception:  # noqa: BLE001
        return None
    if not style_name:
        return None
    lower = style_name.lower()
    if "heading" not in lower and "标题" not in style_name:
        return None
    number_match = re.search(r"(\d+)", style_name)
    if not number_match:
        return 1
    try:
        return max(1, min(6, int(number_match.group(1))))
    except ValueError:
        return 1


def _is_body_text_candidate(text: str) -> bool:
    line = re.sub(r"\s+", "", text)
    if len(line) < 12:
        return False
    if _infer_manual_toc_level_from_text(text) is not None:
        return False
    if _is_reference_heading_text(text):
        return False
    return True


def _is_reference_heading_text(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text.strip().lower())
    return normalized in {"参考文献", "references", "reference"}


def _is_reference_entry_text(paragraph: Any, text: str) -> bool:
    line = text.strip()
    if re.match(r"^\[\d+\]", line):
        return True
    if re.match(r"^\d+[.)、]\s*", line):
        return True
    first_line = getattr(paragraph.paragraph_format, "first_line_indent", None)
    twips = _length_to_twips(first_line)
    if twips is not None and twips < 0:
        return True
    if "," in line and re.search(r"\b(19|20)\d{2}\b", line):
        return True
    return False


def _is_better_reference_profile(candidate: dict[str, Any], existing: dict[str, Any]) -> bool:
    candidate_first = candidate.get("first_line_indent")
    existing_first = existing.get("first_line_indent")
    if isinstance(candidate_first, int) and candidate_first < 0:
        if not isinstance(existing_first, int) or existing_first >= 0:
            return True
    return False


def _is_toc_intro_text(text: str) -> bool:
    compact = re.sub(r"\s+", "", text)
    return any(keyword in compact for keyword in ("目录", "章节安排", "具体安排如下", "拟分为"))


def _infer_manual_toc_level_from_text(text: str) -> int | None:
    line = re.sub(r"\s+", " ", text.strip())
    if not line:
        return None
    if re.match(r"^\s*第[\u96f6\u3007\u4e00-\u9fa5\d]+\s*章", line):
        return 1
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+\s*[.\uFF0E]\s*\d+\s+.+$", line):
        return 3
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+\s+.+$", line):
        return 2
    if re.match(r"^\s*[一二三四五六七八九十]+\s*、", line):
        return 1
    return None


def _iter_body_blocks(doc: Any) -> list[Any]:
    body = doc.element.body
    return [child for child in body.iterchildren() if child.tag in {qn("w:p"), qn("w:tbl")}]


def _paragraph_from_xml(paragraph_xml: Any, doc: Any) -> Any:
    from docx.text.paragraph import Paragraph

    return Paragraph(paragraph_xml, doc)


def _table_from_xml(table_xml: Any, doc: Any) -> Any:
    from docx.table import Table

    return Table(table_xml, doc)


def _clone_paragraph_to_template(source_paragraph: Any, target_doc: Any) -> Any:
    from docx.text.paragraph import Paragraph

    paragraph_xml = deepcopy(source_paragraph._p)
    body = target_doc.element.body
    sectpr = next((child for child in body.iterchildren() if child.tag == qn("w:sectPr")), None)
    if sectpr is not None:
        body.insert(body.index(sectpr), paragraph_xml)
    else:
        body.append(paragraph_xml)
    return Paragraph(paragraph_xml, target_doc)


def _normalize_paragraph_runs(paragraph: Any, run_profile: dict[str, Any] | None) -> None:
    for run in paragraph.runs:
        _clear_run_direct_format(run)
        _apply_run_profile(run, run_profile)


def _clear_run_direct_format(run: Any) -> None:
    font = run.font
    font.name = None
    font.size = None
    font.bold = None
    font.italic = None
    font.underline = None
    font.color.rgb = None
    font.highlight_color = None
    font.all_caps = None
    font.small_caps = None
    font.strike = None
    font.double_strike = None
    font.subscript = None
    font.superscript = None


def _apply_run_profile(run: Any, run_profile: dict[str, Any] | None) -> None:
    if not run_profile:
        return
    font = run.font

    if run_profile.get("name"):
        font.name = run_profile.get("name")
    if Pt is not None and isinstance(run_profile.get("size_pt"), (int, float)):
        font.size = Pt(float(run_profile["size_pt"]))

    font.bold = run_profile.get("bold")
    font.italic = run_profile.get("italic")
    font.underline = run_profile.get("underline")
    font.highlight_color = run_profile.get("highlight_color")

    color_rgb = run_profile.get("color_rgb")
    if color_rgb:
        try:
            from docx.shared import RGBColor

            font.color.rgb = RGBColor.from_string(str(color_rgb))
        except Exception:  # noqa: BLE001
            pass

    if qn is not None:
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        east_asia = run_profile.get("east_asia")
        ascii_name = run_profile.get("ascii")
        hansi_name = run_profile.get("hansi")
        if east_asia:
            r_fonts.set(qn("w:eastAsia"), str(east_asia))
        if ascii_name:
            r_fonts.set(qn("w:ascii"), str(ascii_name))
        if hansi_name:
            r_fonts.set(qn("w:hAnsi"), str(hansi_name))


def _copy_paragraph_to_template(
    source_paragraph: Any,
    target_doc: Any,
    style_map: dict[str, Any],
    profile_map: dict[str, Any],
    run_profile_map: dict[str, Any],
    context: dict[str, Any],
) -> None:
    role, level = _classify_source_paragraph(source_paragraph, context)
    target_style = _resolve_target_style_name(source_paragraph, style_map, role, level)
    run_profile = _resolve_run_profile_for_role(run_profile_map, role, level)

    target_paragraph = _clone_paragraph_to_template(source_paragraph, target_doc)
    if target_style:
        try:
            target_paragraph.style = target_style
        except Exception:  # noqa: BLE001
            pass
    _reset_paragraph_direct_format(target_paragraph)
    _apply_profile_format(target_paragraph, _resolve_profile_for_role(profile_map, role, level))
    _normalize_paragraph_runs(target_paragraph, run_profile)


def _resolve_target_style_name(
    source_paragraph: Any,
    style_map: dict[str, Any],
    role: str,
    level: int | None,
) -> str:
    if role == "reference_entry":
        reference_style = style_map.get("reference")
        if reference_style:
            return str(reference_style)
    if role == "reference_heading":
        reference_heading_style = style_map.get("reference_heading")
        if reference_heading_style:
            return str(reference_heading_style)
    if role == "toc":
        toc_level = level or _resolve_toc_level(source_paragraph)
        if toc_level is not None:
            toc_style = style_map.get("toc", {}).get(toc_level) or style_map.get("toc", {}).get(1)
            if toc_style:
                return str(toc_style)
    if role == "heading":
        heading_level = level or _resolve_heading_level(source_paragraph)
        if heading_level is not None:
            heading_style = style_map.get("heading", {}).get(heading_level)
            if heading_style:
                return str(heading_style)
    return str(style_map.get("normal") or "Normal")


def _classify_source_paragraph(paragraph: Any, context: dict[str, Any]) -> tuple[str, int | None]:
    text = str(getattr(paragraph, "text", "") or "").strip()
    if not text:
        return "normal", None

    if _is_reference_heading_text(text):
        context["in_references"] = True
        context["toc_mode"] = False
        return "reference_heading", None

    if _is_toc_intro_text(text):
        context["toc_mode"] = True
        return "normal", None

    if bool(context.get("toc_mode")):
        toc_level = _infer_manual_toc_level_from_text(text) or _infer_toc_level_from_text(text)
        if toc_level is not None:
            return "toc", toc_level
        context["toc_mode"] = False

    if bool(context.get("in_references")) and _is_reference_entry_text(paragraph, text):
        return "reference_entry", None

    heading_level = _resolve_heading_level(paragraph)
    if heading_level is not None:
        return "heading", heading_level

    toc_level = _resolve_toc_level(paragraph)
    if toc_level is not None:
        return "toc", toc_level

    return "normal", None


def _resolve_heading_level(paragraph: Any) -> int | None:
    try:
        style_name = str(getattr(paragraph.style, "name", "") or "")
    except Exception:  # noqa: BLE001
        return None

    if not style_name:
        text = str(getattr(paragraph, "text", "") or "").strip()
        return _infer_heading_level_from_text(text)

    lower = style_name.lower()
    if "heading" not in lower and "标题" not in style_name:
        text = str(getattr(paragraph, "text", "") or "").strip()
        return _infer_heading_level_from_text(text)

    number_match = re.search(r"(\d+)", style_name)
    if not number_match:
        return 1
    try:
        return max(1, min(6, int(number_match.group(1))))
    except ValueError:
        return 1


def _resolve_toc_level(paragraph: Any) -> int | None:
    try:
        style_name = str(getattr(paragraph.style, "name", "") or "")
    except Exception:  # noqa: BLE001
        style_name = ""

    if style_name:
        lower = style_name.lower()
        if "toc" in lower or "目录" in style_name:
            number_match = re.search(r"(\d+)", style_name)
            if number_match:
                try:
                    return max(1, min(6, int(number_match.group(1))))
                except ValueError:
                    return 1
            return 1

    text = str(getattr(paragraph, "text", "") or "").strip()
    return _infer_toc_level_from_text(text)


def _infer_heading_level_from_text(text: str) -> int | None:
    line = re.sub(r"\s+", " ", text.strip())
    if not line or len(line) > 80:
        return None
    if re.search(r"[。！？!?；;]$", line):
        return None

    if re.match(r"^\s*\u7b2c[\u96f6\u3007\u4e00-\u9fa5\d]+\s*\u7ae0(?:[\s:：.\-]+.+)?$", line):
        return 1
    if re.match(r"^\s*\d+\s*[.\uFF0E]\s*\d+(?:\s*[.\uFF0E]\s*\d+)?\s*(.+)$", line):
        return 2
    if re.match(r"^\s*\d{1,2}(?:\s*[.\uFF0E、)]\s*|\s+).+$", line):
        return 1
    return None


def _infer_toc_level_from_text(text: str) -> int | None:
    line = re.sub(r"\s+", " ", text.strip())
    if not line:
        return None
    if line.lower() in {"contents"} or line == "目录":
        return 1

    toc_match = re.match(r"^(?P<title>.+?)\s*[.·•…]{2,}\s*(?P<page>\d+)\s*$", line)
    if not toc_match:
        return None
    title = toc_match.group("title").strip()
    if re.match(r"^\d+[.\uFF0E]\d+[.\uFF0E]\d+", title):
        return 3
    if re.match(r"^\d+[.\uFF0E]\d+", title):
        return 2
    return 1


def _reset_paragraph_direct_format(paragraph: Any) -> None:
    paragraph.alignment = None
    fmt = paragraph.paragraph_format
    fmt.left_indent = None
    fmt.right_indent = None
    fmt.first_line_indent = None
    fmt.space_before = None
    fmt.space_after = None
    fmt.line_spacing = None
    fmt.line_spacing_rule = None
    fmt.keep_together = None
    fmt.keep_with_next = None
    fmt.page_break_before = None
    fmt.widow_control = None


def _resolve_heading_style_name(paragraph: Any) -> str | None:
    heading_level = _resolve_heading_level(paragraph)
    if heading_level is None:
        return None
    try:
        style_name = str(getattr(paragraph.style, "name", "") or "")
    except Exception:  # noqa: BLE001
        return None
    if style_name:
        return style_name
    return None


def _copy_table_to_template(
    source_table: Any,
    target_doc: Any,
    table_style: str | None,
    paragraph_profile: dict[str, Any] | None,
    run_profile: dict[str, Any] | None,
) -> None:
    table_xml = deepcopy(source_table._tbl)
    body = target_doc.element.body
    sectpr = next((child for child in body.iterchildren() if child.tag == qn("w:sectPr")), None)
    if sectpr is not None:
        body.insert(body.index(sectpr), table_xml)
    else:
        body.append(table_xml)

    target_table = _table_from_xml(table_xml, target_doc)
    if table_style:
        try:
            target_table.style = table_style
        except Exception:  # noqa: BLE001
            pass

    for row in target_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _reset_paragraph_direct_format(paragraph)
                _apply_profile_format(paragraph, paragraph_profile)
                _normalize_paragraph_runs(paragraph, run_profile)


def _build_result(
    operation: str,
    source: Path,
    destination: Path | None,
    status: OperationStatus,
    message: str,
    started: datetime,
    started_at: str,
) -> OperationResult:
    finished = datetime.now()
    finished_at = now_iso()
    return OperationResult(
        operation=operation,
        source=str(source),
        destination=str(destination) if destination else None,
        status=status,
        message=message,
        started_at=started_at,
        finished_at=finished_at,
        duration_ms=duration_ms(started, finished),
    )
