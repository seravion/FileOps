from __future__ import annotations

import json
import re
import shutil
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any
from zipfile import is_zipfile

from .models import OperationResult, OperationStatus
from .utils import duration_ms, ensure_workspace_path, now_iso

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:  # pragma: no cover
    Document = None
    PackageNotFoundError = None
    Table = None
    Paragraph = None
    qn = None

try:
    import pytesseract
    from PIL import Image
except ImportError:  # pragma: no cover
    pytesseract = None
    Image = None

try:
    from pypdf import PdfReader, PdfWriter
except ImportError:  # pragma: no cover
    PdfReader = None
    PdfWriter = None


HEADING_MODES = {"h1", "h2", "h1_h2"}
INPUT_FORMATS = {"auto", "docx", "markdown", "txt", "pdf"}
OUTPUT_FORMATS = {"auto", "docx", "md", "txt", "pdf"}
TEXT_EXTENSIONS = {".md", ".markdown", ".txt"}
PDF_EXTENSIONS = {".pdf"}


def split_documents_by_structure(
    sources: list[Path],
    destination: Path,
    workspace: Path,
    dry_run: bool,
    heading_mode: str,
    include_image_text: bool,
    input_format: str = "auto",
    output_format: str = "auto",
) -> list[OperationResult]:
    if heading_mode not in HEADING_MODES:
        raise ValueError(f"Unsupported heading mode: {heading_mode}")
    if input_format not in INPUT_FORMATS:
        raise ValueError(f"Unsupported input format: {input_format}")
    if output_format not in OUTPUT_FORMATS:
        raise ValueError(f"Unsupported output format: {output_format}")

    destination = destination.resolve(strict=False)
    ensure_workspace_path(destination, workspace)

    if not dry_run:
        destination.mkdir(parents=True, exist_ok=True)

    results: list[OperationResult] = []

    for source in sources:
        started = datetime.now()
        started_at = now_iso()

        try:
            ensure_workspace_path(source, workspace)
            if not source.exists():
                raise FileNotFoundError(f"Source does not exist: {source}")
            if source.is_dir():
                raise IsADirectoryError(f"Document split supports files only: {source}")

            ext = source.suffix.lower()
            if not _matches_input_format(ext, input_format):
                raise ValueError(f"Source format does not match import format setting: {source.name}")

            if ext == ".docx":
                sections = _split_docx(source, heading_mode, include_image_text)
                if not sections:
                    sections = [{"title": "section", "start": 0, "end": 0}]
            elif ext in PDF_EXTENSIONS:
                sections = _split_pdf_document(source, heading_mode, include_image_text)
                if not sections:
                    sections = [{"title": "section", "lines": [""]}]
            elif ext in TEXT_EXTENSIONS:
                sections = _split_text_document(source, heading_mode, include_image_text)
                if not sections:
                    sections = [{"title": "section", "lines": [""]}]
            else:
                raise ValueError(f"Unsupported file type for document split: {source.name}")

            resolved_output = _resolve_output_format(ext, output_format)
            if resolved_output == "pdf" and ext not in PDF_EXTENSIONS:
                raise ValueError("PDF export is currently supported only when the input is PDF.")

            if ext in TEXT_EXTENSIONS:
                sections = _split_line_sections_by_inline_headings(_sanitize_line_sections(sections), heading_mode)
            elif ext in PDF_EXTENSIONS and resolved_output != "pdf":
                sections = _split_line_sections_by_inline_headings(_sanitize_line_sections(sections), heading_mode)

            if dry_run:
                message = f"Would split into {len(sections)} section(s), output format: {resolved_output}."
                results.append(
                    _build_result("doc_split", source, destination, OperationStatus.DRY_RUN, message, started, started_at)
                )
                continue

            if ext == ".docx":
                if resolved_output == "docx":
                    created = _write_docx_sections(source, destination, sections)
                else:
                    created = _write_docx_sections_as_text(
                        source=source,
                        destination=destination,
                        sections=sections,
                        output_ext=resolved_output,
                        include_image_text=include_image_text,
                    )
            elif ext in PDF_EXTENSIONS:
                if resolved_output == "pdf":
                    created = _write_pdf_sections(source, destination, sections)
                elif resolved_output == "docx":
                    created = _write_text_sections_as_docx(source, destination, sections)
                else:
                    created = _write_text_sections(source, destination, sections, output_ext=resolved_output)
            else:
                if resolved_output == "docx":
                    created = _write_text_sections_as_docx(source, destination, sections)
                else:
                    created = _write_text_sections(source, destination, sections, output_ext=resolved_output)

            message = f"Document split completed: {len(created)} file(s) generated."
            results.append(_build_result("doc_split", source, destination, OperationStatus.SUCCESS, message, started, started_at))

        except Exception as exc:  # noqa: BLE001
            results.append(_build_result("doc_split", source, None, OperationStatus.FAILED, str(exc), started, started_at))

    return results


def _matches_input_format(ext: str, input_format: str) -> bool:
    if input_format == "auto":
        return ext == ".docx" or ext in TEXT_EXTENSIONS or ext in PDF_EXTENSIONS
    if input_format == "docx":
        return ext == ".docx"
    if input_format == "markdown":
        return ext in {".md", ".markdown"}
    if input_format == "pdf":
        return ext in PDF_EXTENSIONS
    return ext == ".txt"


def _resolve_output_format(source_ext: str, output_format: str) -> str:
    if output_format != "auto":
        return output_format
    if source_ext == ".docx":
        return "docx"
    if source_ext == ".pdf":
        return "pdf"
    if source_ext in {".md", ".markdown"}:
        return "md"
    return "txt"


def _looks_like_legacy_doc(source: Path) -> bool:
    try:
        with source.open("rb") as stream:
            signature = stream.read(8)
    except OSError:
        return False
    return signature == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _split_docx(source: Path, heading_mode: str, include_image_text: bool) -> list[dict[str, Any]]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    if not is_zipfile(source):
        if _looks_like_legacy_doc(source):
            raise ValueError("Detected legacy DOC or fake .docx. Please re-save as standard .docx in Word/WPS and retry.")
        raise ValueError("Invalid .docx package (file damaged or extension mismatch). Please re-save as .docx in Word/WPS and retry.")

    _ = include_image_text

    try:
        with source.open("rb") as stream:
            doc = Document(stream)

        blocks = _iter_docx_body_blocks(doc)
        if not blocks:
            return [{"title": "section", "start": 0, "end": 0}]

        heading_points: list[tuple[int, str]] = []
        for idx, block in enumerate(blocks):
            if block.tag != qn("w:p"):
                continue
            paragraph = _paragraph_from_xml(block, doc)
            heading_level = _get_docx_heading_level_from_paragraph(paragraph)
            if not _is_heading_boundary(heading_level, heading_mode):
                continue
            title = paragraph.text.strip() or f"section_{len(heading_points) + 1}"
            heading_points.append((idx, title))

        sections: list[dict[str, Any]] = []
        if not heading_points:
            sections.append({"title": "Preface", "start": 0, "end": len(blocks)})
            return _normalize_docx_sections(sections, len(blocks))

        first_start = heading_points[0][0]
        if first_start > 0:
            sections.append({"title": "Preface", "start": 0, "end": first_start})

        for index, (start_idx, title) in enumerate(heading_points):
            end_idx = heading_points[index + 1][0] if index + 1 < len(heading_points) else len(blocks)
            if end_idx > start_idx:
                sections.append({"title": title, "start": start_idx, "end": end_idx})

        return _normalize_docx_sections(sections, len(blocks))

    except Exception as exc:  # noqa: BLE001
        if PackageNotFoundError is not None and isinstance(exc, PackageNotFoundError):
            raise ValueError("Unable to read .docx package. Ensure the file exists and valid.") from exc
        raise ValueError(f".docx parse failed: {exc}") from exc


def _split_text_document(source: Path, heading_mode: str, include_image_text: bool) -> list[dict[str, Any]]:
    lines = source.read_text(encoding="utf-8", errors="ignore").splitlines()

    sections: list[dict[str, Any]] = []
    current_title = "Preface"
    current_lines: list[str] = []

    heading_regex = re.compile(r"^(#{1,6})\s+(.*)$")
    image_regex = re.compile(r"!\[(.*?)\]\((.*?)\)")

    for raw in lines:
        line = raw.rstrip()
        match = heading_regex.match(line)
        if match:
            level = len(match.group(1))
            title = match.group(2).strip() or f"section_{len(sections) + 1}"
            if _is_heading_boundary(level, heading_mode):
                if current_lines:
                    sections.append({"title": current_title, "lines": current_lines[:]})
                    current_lines.clear()
                current_title = title
            current_lines.append(line)
            continue

        current_lines.append(line)

        if include_image_text:
            image_match = image_regex.search(line)
            if image_match:
                alt_text = image_match.group(1).strip()
                if alt_text:
                    current_lines.append(f"[Image Alt] {alt_text}")

    if current_lines:
        sections.append({"title": current_title, "lines": current_lines[:]})

    return sections


def _split_pdf_document(source: Path, heading_mode: str, include_image_text: bool) -> list[dict[str, Any]]:
    if PdfReader is None:
        raise RuntimeError("pypdf is not installed. Please install dependencies first.")

    _ = include_image_text

    try:
        with source.open("rb") as stream:
            reader = PdfReader(stream)
            if bool(getattr(reader, "is_encrypted", False)):
                try:
                    decrypt_result = reader.decrypt("")
                except Exception as exc:  # noqa: BLE001
                    if "cryptography" in str(exc).lower():
                        raise RuntimeError("Cannot read encrypted PDF. Build is missing dependency: cryptography.") from exc
                    raise ValueError(f"Encrypted PDF cannot be read: {exc}") from exc
                if int(decrypt_result) == 0:
                    raise ValueError("PDF is encrypted and requires a password. Please remove protection or provide an unencrypted file.")
            total_pages = len(reader.pages)
            pages_lines: list[list[str]] = []
            for page in reader.pages:
                page_text = page.extract_text() or ""
                normalized = [_normalize_pdf_line(line) for line in page_text.splitlines()]
                pages_lines.append([line for line in normalized if line])

            outline_sections = _build_pdf_sections_from_outline(
                reader=reader,
                pages_lines=pages_lines,
                heading_mode=heading_mode,
                total_pages=total_pages,
            )
            if outline_sections:
                return outline_sections

            header_lines, footer_lines = _detect_repeated_pdf_margin_lines(pages_lines)
            entries: list[dict[str, Any]] = []
            seen_top_header_lines: set[str] = set()
            for page_index, page_lines in enumerate(pages_lines):
                line_count = len(page_lines)
                for line_index, line in enumerate(page_lines):
                    if line_index < 2 and line in header_lines:
                        if line in seen_top_header_lines:
                            continue
                        seen_top_header_lines.add(line)
                    if line_index >= max(0, line_count - 2) and line in footer_lines:
                        continue
                    entries.append({"page": page_index, "line_no": line_index, "line": line})

        if not entries:
            return [{"title": "section", "lines": [""], "start_page": 0, "end_page": max(1, total_pages)}]

        boundaries: list[dict[str, Any]] = []
        for index, entry in enumerate(entries):
            heading_level = _get_pdf_heading_level(str(entry["line"]))
            if not _is_heading_boundary(heading_level, heading_mode):
                continue
            boundaries.append(
                {
                    "index": index,
                    "page": int(entry["page"]),
                    "line_no": int(entry["line_no"]),
                    "title": str(entry["line"]),
                    "level": int(heading_level or 0),
                }
            )

        boundaries = _filter_pdf_boundaries(boundaries, heading_mode, total_pages)

        sections: list[dict[str, Any]] = []

        def _append_section(title: str, start_index: int, end_index: int, start_page: int, end_page: int) -> None:
            if end_index <= start_index:
                return
            lines = [str(entries[idx]["line"]) for idx in range(start_index, end_index)]
            safe_start_page = max(0, start_page)
            safe_end_page = max(safe_start_page + 1, end_page)
            sections.append(
                {
                    "title": title.strip() or f"section_{len(sections) + 1}",
                    "lines": lines,
                    "start_page": safe_start_page,
                    "end_page": safe_end_page,
                }
            )

        if not boundaries:
            _append_section("Preface", 0, len(entries), 0, max(1, total_pages))
            return sections

        first_boundary = boundaries[0]
        first_index = int(first_boundary["index"])
        first_page = int(first_boundary["page"])
        if first_index > 0:
            _append_section("Preface", 0, first_index, 0, max(1, first_page))

        for idx, boundary in enumerate(boundaries):
            start_index = int(boundary["index"])
            title = str(boundary["title"])
            start_page = int(boundary["page"])
            next_index = int(boundaries[idx + 1]["index"]) if idx + 1 < len(boundaries) else len(entries)
            next_page = int(boundaries[idx + 1]["page"]) if idx + 1 < len(boundaries) else total_pages
            _append_section(title, start_index, next_index, start_page, max(start_page + 1, next_page))

        return sections
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f".pdf parse failed: {exc}") from exc


def _build_pdf_sections_from_outline(
    reader: Any,
    pages_lines: list[list[str]],
    heading_mode: str,
    total_pages: int,
) -> list[dict[str, Any]]:
    raw_outline = getattr(reader, "outline", None)
    if not isinstance(raw_outline, list) or not raw_outline:
        return []

    flat_items: list[dict[str, Any]] = []
    _flatten_pdf_outline_items(reader=reader, outline=raw_outline, level=1, output=flat_items)
    if not flat_items:
        return []

    base_level = min(int(item["level"]) for item in flat_items)
    level_set = {int(item["level"]) for item in flat_items}
    if heading_mode == "h1":
        target_levels = {base_level}
    elif heading_mode == "h2":
        h2_level = base_level + 1
        if h2_level not in level_set:
            return []
        target_levels = {h2_level}
    else:
        h2_level = base_level + 1
        target_levels = {base_level}
        if h2_level in level_set:
            target_levels.add(h2_level)

    candidates = [item for item in flat_items if int(item["level"]) in target_levels]
    if not candidates:
        return []

    deduped: list[dict[str, Any]] = []
    seen_pairs: set[tuple[int, str]] = set()
    for item in candidates:
        page = int(item["page"])
        key = (page, _normalize_pdf_heading_key(str(item["title"])))
        if key in seen_pairs:
            continue
        seen_pairs.add(key)
        deduped.append(item)

    if not deduped:
        return []

    sections: list[dict[str, Any]] = []

    def _append_outline_section(title: str, start_page: int, end_page: int) -> None:
        safe_start = max(0, start_page)
        safe_end = max(safe_start + 1, min(end_page, total_pages))
        lines = _collect_pdf_lines_by_page_range(pages_lines, safe_start, safe_end)
        if not lines:
            lines = [""]
        sections.append(
            {
                "title": title.strip() or f"section_{len(sections) + 1}",
                "lines": lines,
                "start_page": safe_start,
                "end_page": safe_end,
            }
        )

    first_page = int(deduped[0]["page"])
    if first_page > 0:
        _append_outline_section("Preface", 0, first_page)

    for idx, item in enumerate(deduped):
        start_page = int(item["page"])
        title = str(item["title"])
        next_page = int(deduped[idx + 1]["page"]) if idx + 1 < len(deduped) else total_pages
        if next_page <= start_page:
            next_page = min(total_pages, start_page + 1)
        _append_outline_section(title, start_page, next_page)

    return sections


def _flatten_pdf_outline_items(reader: Any, outline: list[Any], level: int, output: list[dict[str, Any]]) -> None:
    last_item: Any = None
    for node in outline:
        if isinstance(node, list):
            if last_item is not None:
                _flatten_pdf_outline_items(reader, node, level + 1, output)
            continue

        title = str(getattr(node, "title", "") or "").strip()
        page_number = _get_pdf_outline_page_number(reader, node)
        if title and page_number is not None:
            output.append({"title": title, "page": page_number, "level": level})
            last_item = node


def _get_pdf_outline_page_number(reader: Any, node: Any) -> int | None:
    try:
        page = reader.get_destination_page_number(node)
    except Exception:  # noqa: BLE001
        return None
    if page is None:
        return None
    try:
        page_int = int(page)
    except Exception:  # noqa: BLE001
        return None
    return page_int if page_int >= 0 else None


def _collect_pdf_lines_by_page_range(pages_lines: list[list[str]], start_page: int, end_page: int) -> list[str]:
    lines: list[str] = []
    for page_index in range(start_page, min(end_page, len(pages_lines))):
        page_content = pages_lines[page_index]
        if page_content:
            lines.extend(page_content)
        if page_index < end_page - 1 and lines and lines[-1] != "":
            lines.append("")
    while lines and not str(lines[-1]).strip():
        lines.pop()
    return lines


def _normalize_pdf_line(line: str) -> str:
    compact = re.sub(r"\s+", " ", line.replace("\u3000", " ").strip())
    if not compact:
        return ""
    compact = compact.strip("\u00b7\u2022")
    if re.fullmatch(r"(?i)page\s*\d+", compact):
        return ""
    if re.fullmatch(r"^\u7b2c?\d+\s*\u9875$", compact):
        return ""
    if re.fullmatch(r"-?\d{1,4}-?", compact):
        return ""
    return compact


def _get_pdf_heading_level(line: str) -> int | None:
    if not line:
        return None

    markdown_match = re.match(r"^(#{1,6})\s+(.+)$", line)
    if markdown_match:
        return len(markdown_match.group(1))

    zh_h1 = re.match(
        r"^\u7b2c[\u96f6\u3007\u4e00\u4e8c\u4e24\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u5343\d]+\u7ae0(?:[:\uFF1A.\-\s]*.*)?$",
        line,
    )
    if zh_h1 and _is_probable_pdf_heading_tail(line.split("\u7ae0", 1)[1]):
        return 1

    zh_h2 = re.match(
        r"^\u7b2c[\u96f6\u3007\u4e00\u4e8c\u4e24\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\u5343\d]+\u8282(?:[:\uFF1A.\-\s]*.*)?$",
        line,
    )
    if zh_h2 and _is_probable_pdf_heading_tail(line.split("\u8282", 1)[1]):
        return 2

    en_h1 = re.match(r"^(chapter|appendix)\s*([a-z\d]+)(?:[\s:\uFF1A.\-]*)(.*)$", line, flags=re.IGNORECASE)
    if en_h1:
        tail = en_h1.group(3).strip()
        if not tail or _is_probable_pdf_heading_tail(tail):
            return 1

    multi_level_number = re.match(r"^(\d+(?:[.\uFF0E]\d+)+)(?:[\u3001.\uFF0E)\uFF09]?\s*)(.+)$", line)
    if multi_level_number:
        marker = multi_level_number.group(1).replace("\uFF0E", ".")
        tail = multi_level_number.group(2).strip()
        marker_parts = [int(item) for item in marker.split(".") if item.isdigit()]
        if marker_parts and all(0 < item <= 99 for item in marker_parts) and _is_probable_pdf_heading_tail(tail):
            return 2

    single_level_number = re.match(r"^(\d{1,2})(?:(?:[\u3001.\uFF0E)\uFF09]\s*)|(?:\s+))(.+)$", line)
    if single_level_number:
        tail = single_level_number.group(2).strip()
        if _is_probable_pdf_heading_tail(tail):
            return 1

    upper_title = re.match(r"^[A-Z][A-Z\s\-:]{3,64}$", line)
    if upper_title and len(line.split()) <= 8:
        return 1

    return None


def _is_probable_pdf_heading_tail(text: str) -> bool:
    cleaned = text.strip(":\uFF1A.-_\u3001)\uFF09 ")
    if not cleaned:
        return True
    if len(cleaned) < 2 or len(cleaned) > 80:
        return False
    if re.search(r"[=+\-*/<>^_{}\[\]|~`$]", cleaned):
        return False
    if re.search(r"[\u0391-\u03C9\u2200-\u22FF]", cleaned):
        return False
    if re.search(r"\b(for|while|if|else|return|end|do|then|repeat|until)\b", cleaned, flags=re.IGNORECASE):
        return False
    if not re.search(r"[\u4e00-\u9fffA-Za-z]", cleaned):
        return False
    if re.search(r"[\u3002\uFF01\uFF1F!?;\uFF1B]$", cleaned):
        return False
    return True


def _detect_repeated_pdf_margin_lines(pages_lines: list[list[str]]) -> tuple[set[str], set[str]]:
    if not pages_lines:
        return set(), set()

    header_counter: Counter[str] = Counter()
    footer_counter: Counter[str] = Counter()
    pages_with_text = 0

    for lines in pages_lines:
        if not lines:
            continue
        pages_with_text += 1
        top_candidates = {line for line in lines[:2] if len(line) <= 80}
        bottom_candidates = {line for line in lines[-2:] if len(line) <= 80}
        header_counter.update(top_candidates)
        footer_counter.update(bottom_candidates)

    if pages_with_text < 3:
        return set(), set()

    threshold = max(3, int(pages_with_text * 0.4))
    headers = {text for text, count in header_counter.items() if count >= threshold}
    footers = {text for text, count in footer_counter.items() if count >= threshold}
    return headers, footers


def _filter_pdf_boundaries(boundaries: list[dict[str, Any]], heading_mode: str, total_pages: int) -> list[dict[str, Any]]:
    if not boundaries:
        return boundaries

    title_counter: Counter[str] = Counter(_normalize_pdf_heading_key(str(item["title"])) for item in boundaries)
    repeated_threshold = max(3, int(max(1, total_pages) * 0.3))
    seen_repeated_titles: set[str] = set()

    filtered: list[dict[str, Any]] = []
    for boundary in boundaries:
        key = _normalize_pdf_heading_key(str(boundary["title"]))
        if title_counter[key] >= repeated_threshold:
            if key in seen_repeated_titles:
                continue
            seen_repeated_titles.add(key)
        filtered.append(boundary)

    if not filtered:
        return boundaries[:1]

    pruned: list[dict[str, Any]] = [filtered[0]]
    min_line_gap = 8 if heading_mode == "h1" else 2
    for boundary in filtered[1:]:
        prev = pruned[-1]
        page_gap = int(boundary["page"]) - int(prev["page"])
        line_gap = int(boundary["index"]) - int(prev["index"])
        if _normalize_pdf_heading_key(str(boundary["title"])) == _normalize_pdf_heading_key(str(prev["title"])) and page_gap <= 2:
            continue
        if page_gap == 0 and line_gap < min_line_gap:
            continue
        pruned.append(boundary)

    return pruned


def _normalize_pdf_heading_key(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def _iter_docx_body_blocks(doc: Any) -> list[Any]:
    if qn is None:
        return []
    body = doc.element.body
    return [child for child in body.iterchildren() if child.tag in {qn("w:p"), qn("w:tbl")}]


def _paragraph_from_xml(paragraph_xml: Any, doc: Any) -> Any:
    if Paragraph is None:
        raise RuntimeError("python-docx paragraph support is unavailable.")
    return Paragraph(paragraph_xml, doc)


def _table_from_xml(table_xml: Any, doc: Any) -> Any:
    if Table is None:
        raise RuntimeError("python-docx table support is unavailable.")
    return Table(table_xml, doc)


def _safe_docx_style_name(paragraph: Any) -> str:
    try:
        style = paragraph.style
    except Exception:  # noqa: BLE001
        return ""
    return str(getattr(style, "name", "") or "")


def _extract_docx_image_lines(paragraph: Any, include_image_text: bool) -> list[str]:
    if not include_image_text:
        return []

    image_lines: list[str] = []
    if qn is None:
        return image_lines

    try:
        blips = paragraph._element.xpath(".//a:blip")
    except Exception:  # noqa: BLE001
        return image_lines
    for blip in blips:
        embed_id = blip.get(qn("r:embed"))
        if not embed_id:
            continue

        try:
            image_part = paragraph.part.related_parts.get(embed_id)
        except Exception:  # noqa: BLE001
            continue
        if image_part is None:
            continue

        ocr_text = _ocr_image_blob(image_part.blob)
        if ocr_text:
            image_lines.append(f"[Image Text] {ocr_text}")
        else:
            image_lines.append("[Image Text] Not recognized")

    return image_lines


def _ocr_image_blob(blob: bytes) -> str:
    if pytesseract is None or Image is None:
        return ""

    try:
        from io import BytesIO

        image = Image.open(BytesIO(blob))
        text = pytesseract.image_to_string(image, lang="chi_sim+eng")
        return text.strip()
    except Exception:  # noqa: BLE001
        return ""


def _get_docx_heading_level(style_name: str) -> int | None:
    if not style_name:
        return None

    lower = style_name.lower()
    if "heading" not in lower:
        return None

    nums = re.findall(r"\d+", lower)
    if not nums:
        return 1
    return int(nums[0])


def _get_docx_heading_level_from_paragraph(paragraph: Any) -> int | None:
    style_level = _get_docx_heading_level(_safe_docx_style_name(paragraph))
    if style_level is not None:
        return style_level
    return _infer_docx_heading_level_from_text(str(getattr(paragraph, "text", "") or ""))


def _infer_docx_heading_level_from_text(text: str) -> int | None:
    line = re.sub(r"\s+", " ", text.strip())
    if not line or len(line) > 80:
        return None
    if re.search(r"[。！？!?；;]$", line):
        return None

    if re.match(r"^\s*\u7b2c[\u96f6\u3007\u4e00-\u9fa5\d]+\s*\u7ae0(?:[\s:：.\-]+.+)?$", line):
        return 1

    numbered_h2 = re.match(r"^\s*(\d+)\s*[.\uFF0E]\s*(\d+)(?:\s*[.\uFF0E]\s*\d+)?\s*(.+)$", line)
    if numbered_h2:
        tail = numbered_h2.group(3).strip()
        if 1 <= len(tail) <= 60 and not re.search(r"[。！？!?；;]$", tail):
            return 2

    numbered_h1 = re.match(r"^\s*(\d{1,2})(?:\s*[.\uFF0E、)]\s*|\s+)(.+)$", line)
    if numbered_h1:
        tail = numbered_h1.group(2).strip()
        if 1 <= len(tail) <= 50 and not re.search(r"[，,]", tail):
            return 1

    return None


def _is_heading_boundary(level: int | None, mode: str) -> bool:
    if level is None:
        return False
    if mode == "h1":
        return level == 1
    if mode == "h2":
        return level == 2
    return level in {1, 2}


def _normalize_docx_sections(sections: list[dict[str, Any]], total_blocks: int) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    cursor = 0

    for idx, section in enumerate(sections, start=1):
        title = str(section.get("title", "section")).strip() or f"section_{idx}"

        try:
            start = int(section.get("start", cursor))
        except Exception:  # noqa: BLE001
            start = cursor
        try:
            end = int(section.get("end", start))
        except Exception:  # noqa: BLE001
            end = start

        start = max(cursor, min(max(0, start), total_blocks))
        end = max(start, min(max(0, end), total_blocks))

        if end <= start:
            continue

        normalized.append({"title": title, "start": start, "end": end})
        cursor = end

    if not normalized and total_blocks > 0:
        return [{"title": "Preface", "start": 0, "end": total_blocks}]
    return normalized


def _sanitize_line_sections(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not sections:
        return sections

    cleaned: list[dict[str, Any]] = []
    for idx, section in enumerate(sections):
        title = str(section.get("title", "section")).strip() or f"section_{idx + 1}"
        lines = [str(item) for item in section.get("lines", [])]

        while len(lines) >= 2 and _normalize_section_heading(lines[0]) == _normalize_section_heading(lines[1]):
            lines.pop(0)

        if idx > 0:
            prev_title = str(sections[idx - 1].get("title", "") or "")
            while lines and _normalize_section_heading(lines[0]) == _normalize_section_heading(prev_title):
                lines.pop(0)

        if idx + 1 < len(sections):
            next_title = str(sections[idx + 1].get("title", "") or "")
            next_key = _normalize_section_heading(next_title)
            while lines and _normalize_section_heading(lines[-1]) == next_key:
                lines.pop()
            while lines and not lines[-1].strip():
                lines.pop()

        while lines and not lines[-1].strip():
            lines.pop()

        cleaned.append({**section, "title": title, "lines": lines or [""]})

    return cleaned


def _split_line_sections_by_inline_headings(sections: list[dict[str, Any]], heading_mode: str) -> list[dict[str, Any]]:
    if not sections:
        return sections

    split_sections: list[dict[str, Any]] = []
    for idx, section in enumerate(sections, start=1):
        title = str(section.get("title", "section")).strip() or f"section_{idx}"
        lines = [str(item) for item in section.get("lines", [])]
        if not lines:
            split_sections.append({**section, "title": title, "lines": [""]})
            continue

        current_title = title
        current_lines: list[str] = []
        for line_idx, line in enumerate(lines):
            heading_level = _get_inline_heading_level(line)
            line_title = _normalize_line_heading_title(line)
            should_split = (
                line_idx > 0
                and _is_heading_boundary(heading_level, heading_mode)
                and line_title
                and _normalize_section_heading(current_title) != _normalize_section_heading(line_title)
                and current_lines
            )
            if should_split:
                split_sections.append({**section, "title": current_title, "lines": current_lines[:]})
                current_title = line_title
                current_lines = [line]
            else:
                current_lines.append(line)

        if current_lines:
            split_sections.append({**section, "title": current_title, "lines": current_lines[:]})

    return _sanitize_line_sections(split_sections)


def _get_inline_heading_level(line: str) -> int | None:
    stripped = str(line).strip()
    if not stripped:
        return None

    markdown_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if markdown_match:
        return len(markdown_match.group(1))

    pdf_like = _get_pdf_heading_level(stripped)
    if pdf_like is not None:
        return pdf_like

    return _infer_docx_heading_level_from_text(stripped)


def _normalize_line_heading_title(line: str) -> str:
    stripped = str(line).strip()
    markdown_match = re.match(r"^#{1,6}\s+(.+)$", stripped)
    if markdown_match:
        return markdown_match.group(1).strip()
    return stripped


def _normalize_section_heading(text: str) -> str:
    compact = re.sub(r"\s+", " ", str(text).strip())
    compact = re.sub(r"^#{1,6}\s*", "", compact)
    return compact.lower()


def _write_docx_sections(source: Path, destination: Path, sections: list[dict[str, Any]]) -> list[Path]:
    destination.mkdir(parents=True, exist_ok=True)

    created_files: list[Path] = []
    index_payload: list[dict[str, Any]] = []

    for idx, section in enumerate(sections, start=1):
        title = str(section.get("title", "section")).strip() or f"section_{idx}"
        safe_title = _safe_filename(title)
        output_path = destination / f"{source.stem}_{idx:03d}_{safe_title}.docx"

        shutil.copy2(source, output_path)
        start = int(section.get("start", 0))
        end = int(section.get("end", 0))
        _prune_docx_file(output_path, start, end)

        created_files.append(output_path)
        index_payload.append({"index": idx, "title": title, "file": output_path.name, "block_count": max(0, end - start)})

    index_file = destination / f"{source.stem}_split_index.json"
    index_file.write_text(json.dumps(index_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return created_files


def _write_docx_sections_as_text(
    source: Path,
    destination: Path,
    sections: list[dict[str, Any]],
    output_ext: str,
    include_image_text: bool,
) -> list[Path]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    destination.mkdir(parents=True, exist_ok=True)

    doc = Document(str(source))
    blocks = _iter_docx_body_blocks(doc)

    created_files: list[Path] = []
    index_payload: list[dict[str, Any]] = []

    markdown = output_ext == "md"
    suffix = ".md" if markdown else ".txt"

    for idx, section in enumerate(sections, start=1):
        title = str(section.get("title", "section")).strip() or f"section_{idx}"
        safe_title = _safe_filename(title)
        output_path = destination / f"{source.stem}_{idx:03d}_{safe_title}{suffix}"

        start = int(section.get("start", 0))
        end = int(section.get("end", 0))
        lines = _extract_docx_section_lines(doc, blocks, start, end, include_image_text, markdown)

        text_content = "\n".join(lines).strip() + "\n"
        output_path.write_text(text_content, encoding="utf-8")

        created_files.append(output_path)
        index_payload.append({"index": idx, "title": title, "file": output_path.name, "line_count": len(lines)})

    index_file = destination / f"{source.stem}_split_index.json"
    index_file.write_text(json.dumps(index_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return created_files


def _write_pdf_sections(source: Path, destination: Path, sections: list[dict[str, Any]]) -> list[Path]:
    if PdfReader is None or PdfWriter is None:
        raise RuntimeError("pypdf is not installed. Please install dependencies first.")

    destination.mkdir(parents=True, exist_ok=True)

    with source.open("rb") as stream:
        reader = PdfReader(stream)
        total_pages = len(reader.pages)

        if total_pages == 0:
            raise ValueError("PDF contains no pages.")

        created_files: list[Path] = []
        index_payload: list[dict[str, Any]] = []

        for idx, section in enumerate(sections, start=1):
            title = str(section.get("title", "section")).strip() or f"section_{idx}"
            safe_title = _safe_filename(title)
            output_path = destination / f"{source.stem}_{idx:03d}_{safe_title}.pdf"

            start_page = int(section.get("start_page", 0))
            end_page = int(section.get("end_page", start_page + 1))

            start_page = max(0, min(start_page, total_pages - 1))
            end_page = max(start_page + 1, min(end_page, total_pages))

            writer = PdfWriter()
            for page_index in range(start_page, end_page):
                writer.add_page(reader.pages[page_index])

            with output_path.open("wb") as output_stream:
                writer.write(output_stream)

            created_files.append(output_path)
            index_payload.append(
                {
                    "index": idx,
                    "title": title,
                    "file": output_path.name,
                    "start_page": start_page + 1,
                    "end_page": end_page,
                    "page_count": end_page - start_page,
                }
            )

    index_file = destination / f"{source.stem}_split_index.json"
    index_file.write_text(json.dumps(index_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return created_files


def _extract_docx_section_lines(
    doc: Any,
    blocks: list[Any],
    start: int,
    end: int,
    include_image_text: bool,
    markdown: bool,
) -> list[str]:
    lines: list[str] = []

    for idx, block in enumerate(blocks):
        if idx < start or idx >= end:
            continue

        if block.tag == qn("w:p"):
            paragraph = _paragraph_from_xml(block, doc)
            text = paragraph.text.strip()
            if text:
                if markdown:
                    level = _get_docx_heading_level(_safe_docx_style_name(paragraph))
                    if level is not None:
                        text = f"{'#' * min(level, 6)} {text}"
                lines.append(text)
            lines.extend(_extract_docx_image_lines(paragraph, include_image_text))
            continue

        if block.tag == qn("w:tbl"):
            table = _table_from_xml(block, doc)
            table_lines = _table_to_lines(table, markdown)
            if table_lines:
                lines.extend(table_lines)
                lines.append("")

    while lines and not lines[-1].strip():
        lines.pop()

    return lines


def _table_to_lines(table: Any, markdown: bool) -> list[str]:
    rows: list[list[str]] = []

    for row in table.rows:
        rows.append([cell.text.replace("\n", " ").strip() for cell in row.cells])

    if not rows:
        return []

    cols = max(len(row) for row in rows)
    norm_rows = [row + [""] * (cols - len(row)) for row in rows]

    if markdown:
        header = norm_rows[0]
        lines = [f"| {' | '.join(header)} |", f"| {' | '.join(['---'] * cols)} |"]
        lines.extend(f"| {' | '.join(row)} |" for row in norm_rows[1:])
        return lines

    return ["\t".join(row).rstrip() for row in norm_rows]


def _prune_docx_file(path: Path, start: int, end: int) -> None:
    doc = Document(str(path))
    blocks = _iter_docx_body_blocks(doc)
    for idx, block in enumerate(blocks):
        if start <= idx < end:
            continue
        block.getparent().remove(block)

    if not _iter_docx_body_blocks(doc):
        doc.add_paragraph("")

    doc.save(str(path))


def _write_text_sections(source: Path, destination: Path, sections: list[dict[str, Any]], output_ext: str = "txt") -> list[Path]:
    destination.mkdir(parents=True, exist_ok=True)

    created_files: list[Path] = []
    index_payload: list[dict[str, Any]] = []
    suffix = ".md" if output_ext == "md" else ".txt"

    for idx, section in enumerate(sections, start=1):
        title = str(section.get("title", "section")).strip() or f"section_{idx}"
        lines = [str(item) for item in section.get("lines", [])]

        safe_title = _safe_filename(title)
        output_path = destination / f"{source.stem}_{idx:03d}_{safe_title}{suffix}"
        text_content = "\n".join(lines).strip() + "\n"
        output_path.write_text(text_content, encoding="utf-8")

        created_files.append(output_path)
        index_payload.append({"index": idx, "title": title, "file": output_path.name, "line_count": len(lines)})

    index_file = destination / f"{source.stem}_split_index.json"
    index_file.write_text(json.dumps(index_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return created_files


def _write_text_sections_as_docx(source: Path, destination: Path, sections: list[dict[str, Any]]) -> list[Path]:
    if Document is None:
        raise RuntimeError("python-docx is not installed. Please install dependencies first.")

    destination.mkdir(parents=True, exist_ok=True)

    created_files: list[Path] = []
    index_payload: list[dict[str, Any]] = []

    for idx, section in enumerate(sections, start=1):
        title = str(section.get("title", "section")).strip() or f"section_{idx}"
        lines = [str(item) for item in section.get("lines", [])]

        safe_title = _safe_filename(title)
        output_path = destination / f"{source.stem}_{idx:03d}_{safe_title}.docx"

        doc = Document()
        _append_text_lines_to_doc(doc, lines)
        if not doc.paragraphs and not doc.tables:
            doc.add_paragraph("")
        doc.save(str(output_path))

        created_files.append(output_path)
        index_payload.append({"index": idx, "title": title, "file": output_path.name, "line_count": len(lines)})

    index_file = destination / f"{source.stem}_split_index.json"
    index_file.write_text(json.dumps(index_payload, ensure_ascii=False, indent=2), encoding="utf-8")

    return created_files


def _append_text_lines_to_doc(doc: Any, lines: list[str]) -> None:
    heading_regex = re.compile(r"^(#{1,6})\s+(.*)$")

    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.rstrip()

        if _looks_like_markdown_table_row(line):
            table_lines = [line]
            index += 1
            while index < len(lines) and _looks_like_markdown_table_row(lines[index].rstrip()):
                table_lines.append(lines[index].rstrip())
                index += 1
            _append_markdown_table_to_doc(doc, table_lines)
            continue

        heading_match = heading_regex.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if text:
                doc.add_heading(text, level=level)
            else:
                doc.add_paragraph(line)
        elif line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
        index += 1


def _looks_like_markdown_table_row(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2


def _append_markdown_table_to_doc(doc: Any, table_lines: list[str]) -> None:
    rows = [_parse_markdown_table_row(item) for item in table_lines]
    if not rows:
        return

    if len(rows) >= 2 and _is_markdown_separator_row(rows[1]):
        rows = [rows[0], *rows[2:]]

    if not rows:
        return

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)

    for row_idx, row in enumerate(rows):
        norm = row + [""] * (cols - len(row))
        for col_idx, value in enumerate(norm):
            table.cell(row_idx, col_idx).text = value


def _parse_markdown_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _is_markdown_separator_row(row: list[str]) -> bool:
    for cell in row:
        normalized = cell.replace(" ", "")
        if not normalized:
            continue
        if not re.fullmatch(r":?-{3,}:?", normalized):
            return False
    return True


def _safe_filename(value: str) -> str:
    cleaned = re.sub(r"[\\/:*?\"<>|\s]+", "_", value).strip("_")
    return cleaned[:40] or "section"


def _build_result(
    operation: str,
    source: Path,
    destination: Path | None,
    status: OperationStatus,
    message: str,
    started: datetime,
    started_at: str,
) -> OperationResult:
    finished = datetime.now()
    finished_at = now_iso()

    return OperationResult(
        operation=operation,
        source=str(source),
        destination=str(destination) if destination else None,
        status=status,
        message=message,
        started_at=started_at,
        finished_at=finished_at,
        duration_ms=duration_ms(started, finished),
    )

