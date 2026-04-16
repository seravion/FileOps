from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

import pytest

from docx import Document

from fileops.document_split import split_documents_by_structure
from fileops.models import OperationStatus
from fileops.operations import CommonOptions, copy_items, delete_items, move_items, rename_items, split_items
import fileops.operations as operations_module


def test_copy_dry_run_does_not_modify_fs() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "source.txt"
        src.write_text("hello", encoding="utf-8")

        options = CommonOptions(workspace=root, dry_run=True, overwrite="never")
        results = copy_items([src], root / "target.txt", options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.DRY_RUN
        assert not (root / "target.txt").exists()


def test_move_with_auto_rename_policy() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "to-move.txt"
        src.write_text("hello", encoding="utf-8")
        existing = root / "dest.txt"
        existing.write_text("existing", encoding="utf-8")

        options = CommonOptions(workspace=root, dry_run=False, overwrite="rename")
        results = move_items([src], root / "dest.txt", options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (root / "dest_1.txt").exists()
        assert existing.exists()


def test_rename_pattern_applies_index() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.txt"
        src.write_text("hello", encoding="utf-8")

        options = CommonOptions(workspace=root, dry_run=False, overwrite="never")
        results = rename_items([src], pattern="{stem}_{index}{ext}", start_index=7, options=options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (root / "sample_7.txt").exists()


def test_delete_hard_removes_file() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "to-delete.txt"
        src.write_text("bye", encoding="utf-8")

        results = delete_items([src], workspace=root, dry_run=False, use_trash=False)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert not src.exists()


def test_split_items_by_size() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "large.bin"
        src.write_bytes(b"abcdefghijklmnopqrstuvwxyz")

        options = CommonOptions(workspace=root, dry_run=False, overwrite="never")
        out_dir = root / "chunks"
        results = split_items([src], destination=out_dir, chunk_size_mb=0.00001, options=options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (out_dir / "large.part001.bin").exists()
        assert (out_dir / "large.part002.bin").exists()
        assert (out_dir / "large.part003.bin").exists()


def test_split_pdf_by_size_outputs_openable_pdfs() -> None:
    pypdf = pytest.importorskip("pypdf")
    PdfReader = pypdf.PdfReader
    PdfWriter = pypdf.PdfWriter

    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "paper.pdf"

        writer = PdfWriter()
        for _ in range(6):
            writer.add_blank_page(width=595, height=842)
        with src.open("wb") as stream:
            writer.write(stream)

        options = CommonOptions(workspace=root, dry_run=False, overwrite="never")
        out_dir = root / "chunks"
        results = split_items([src], destination=out_dir, chunk_size_mb=0.0001, options=options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        parts = sorted(out_dir.glob("paper.part*.pdf"))
        assert len(parts) >= 2

        total_pages = 0
        for part in parts:
            part_reader = PdfReader(str(part))
            assert len(part_reader.pages) >= 1
            total_pages += len(part_reader.pages)

        assert total_pages == 6


def test_split_docx_by_size_outputs_openable_docx_parts() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "paper.docx"

        doc = Document()
        for idx in range(120):
            doc.add_paragraph(f"Paragraph {idx} - " + ("content " * 20))
        doc.save(str(src))

        options = CommonOptions(workspace=root, dry_run=False, overwrite="never")
        out_dir = root / "chunks"
        chunk_size_mb = max(0.0002, (src.stat().st_size / 2) / (1024 * 1024))
        results = split_items([src], destination=out_dir, chunk_size_mb=chunk_size_mb, options=options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        parts = sorted(out_dir.glob("paper.part*.docx"))
        assert len(parts) >= 2

        total_paragraphs = 0
        for part in parts:
            split_doc = Document(str(part))
            assert len(split_doc.paragraphs) >= 1
            total_paragraphs += sum(1 for p in split_doc.paragraphs if p.text.strip())

        assert total_paragraphs >= 120



def test_split_pdf_by_size_uses_source_size_target_when_estimation_is_small(monkeypatch) -> None:
    pypdf = pytest.importorskip("pypdf")
    PdfWriter = pypdf.PdfWriter

    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "paper.pdf"

        writer = PdfWriter()
        for _ in range(6):
            writer.add_blank_page(width=595, height=842)
        with src.open("wb") as stream:
            writer.write(stream)

        monkeypatch.setattr(operations_module, "_estimate_pdf_size_for_pages", lambda *_args, **_kwargs: 1)

        file_size = src.stat().st_size
        chunk_size_mb = max(0.000001, (file_size - 1) / (1024 * 1024))

        options = CommonOptions(workspace=root, dry_run=False, overwrite="never")
        out_dir = root / "chunks2"
        results = split_items([src], destination=out_dir, chunk_size_mb=chunk_size_mb, options=options)

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        parts = sorted(out_dir.glob("paper.part*.pdf"))
        assert len(parts) == 2


def test_doc_split_markdown_by_heading() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.md"
        src.write_text(
            "# Title One\n"
            "Intro line\n"
            "![diagram](img.png)\n"
            "## Sub Two\n"
            "Second section text\n",
            encoding="utf-8",
        )

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=True,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (out_dir / "sample_split_index.json").exists()

        produced = sorted(out_dir.glob("sample_*.md"))
        assert len(produced) >= 2
        assert "[Image Alt] diagram" in produced[0].read_text(encoding="utf-8")


def test_doc_split_text_sanitizes_cross_section_heading_overlap(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.txt"
        src.write_text("placeholder", encoding="utf-8")

        import fileops.document_split as document_split

        monkeypatch.setattr(
            document_split,
            "_split_text_document",
            lambda *_args, **_kwargs: [
                {"title": "第一章 绪论", "lines": ["第一章 绪论", "绪论正文", "第二章 方法"]},
                {"title": "第二章 方法", "lines": ["第二章 方法", "方法正文"]},
            ],
        )

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
            input_format="txt",
            output_format="txt",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        produced = sorted(out_dir.glob("sample_*.txt"))
        assert len(produced) == 2
        first_text = produced[0].read_text(encoding="utf-8")
        second_text = produced[1].read_text(encoding="utf-8")
        assert "第二章 方法" not in first_text
        assert "方法正文" not in first_text
        assert "第二章 方法" in second_text
        assert "方法正文" in second_text


class _BrokenStyleParagraph:
    @property
    def style(self) -> object:
        raise KeyError("missing style")


class _NormalStyleParagraph:
    class _Style:
        name = "Heading 1"

    @property
    def style(self) -> object:
        return self._Style()


def test_safe_docx_style_name_handles_style_lookup_error() -> None:
    from fileops.document_split import _safe_docx_style_name

    assert _safe_docx_style_name(_BrokenStyleParagraph()) == ""
    assert _safe_docx_style_name(_NormalStyleParagraph()) == "Heading 1"

def test_doc_split_invalid_docx_reports_friendly_error() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "not_a_real_docx.docx"
        src.write_text("this is plain text, not docx package", encoding="utf-8")

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.FAILED
        assert ".docx" in results[0].message
        assert "Word/WPS" in results[0].message


def test_doc_split_docx_keeps_docx_and_tables() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.docx"

        doc = Document()
        doc.add_paragraph("封面段落")
        h1 = doc.add_paragraph("第一章")
        h1.style = "Heading 1"
        doc.add_paragraph("第一章正文")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        h2 = doc.add_paragraph("第一章-小节")
        h2.style = "Heading 2"
        doc.add_paragraph("小节正文")
        doc.save(src)

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (out_dir / "sample_split_index.json").exists()

        produced_docx = sorted(out_dir.glob("sample_*.docx"))
        assert len(produced_docx) >= 2
        assert not list(out_dir.glob("sample_*.txt"))

        split_docs = [Document(str(path)) for path in produced_docx]
        assert any("第一章正文" in "\n".join(p.text for p in split_doc.paragraphs) for split_doc in split_docs)
        assert any(len(split_doc.tables) > 0 for split_doc in split_docs)


def test_doc_split_docx_numeric_headings_without_heading_style() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.docx"

        doc = Document()
        doc.add_paragraph("第一章 绪论")
        doc.add_paragraph("章节导语")
        doc.add_paragraph("1.1 研究背景与意义")
        doc.add_paragraph("这是1.1内容")
        doc.add_paragraph("1.2 国内外研究现状")
        doc.add_paragraph("这是1.2内容")
        doc.save(src)

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        produced_docx = sorted(out_dir.glob("sample_*.docx"))
        assert len(produced_docx) >= 3

        section_11 = [path for path in produced_docx if "1.1_研究背景与意义" in path.name]
        section_12 = [path for path in produced_docx if "1.2_国内外研究现状" in path.name]
        assert section_11 and section_12

        doc_11 = Document(str(section_11[0]))
        doc_12 = Document(str(section_12[0]))
        text_11 = "\n".join(p.text for p in doc_11.paragraphs if p.text.strip())
        text_12 = "\n".join(p.text for p in doc_12.paragraphs if p.text.strip())
        assert "这是1.1内容" in text_11
        assert "1.2 国内外研究现状" not in text_11
        assert "这是1.2内容" not in text_11
        assert "1.2 国内外研究现状" in text_12
        assert "这是1.2内容" in text_12


def test_doc_split_markdown_export_docx() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.md"
        src.write_text(
            "# Title One\n"
            "Body text\n"
            "| Col1 | Col2 |\n"
            "| --- | --- |\n"
            "| A | B |\n",
            encoding="utf-8",
        )

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
            input_format="markdown",
            output_format="docx",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        produced_docx = sorted(out_dir.glob("sample_*.docx"))
        assert produced_docx

        first_doc = Document(str(produced_docx[0]))
        text_content = "\n".join(paragraph.text for paragraph in first_doc.paragraphs)
        assert "Title One" in text_content


def test_doc_split_docx_export_markdown() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.docx"

        doc = Document()
        heading = doc.add_paragraph("Chapter One")
        heading.style = "Heading 1"
        doc.add_paragraph("Body paragraph")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        doc.save(src)

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
            output_format="md",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        produced_md = sorted(out_dir.glob("sample_*.md"))
        assert produced_md

        merged_text = "\n".join(path.read_text(encoding="utf-8") for path in produced_md)
        assert "# Chapter One" in merged_text
        assert "| A1 | B1 |" in merged_text


def test_doc_split_import_format_mismatch() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.txt"
        src.write_text("plain text", encoding="utf-8")

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="docx",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.FAILED
        assert "import format" in results[0].message.lower()


def test_doc_split_pdf_by_heading_with_fake_reader(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "paper.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        class _FakePdfPage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class _FakePdfReader:
            def __init__(self, _stream: object) -> None:
                self.pages = [
                    _FakePdfPage("封面\n第一章 绪论\n这里是第一章内容\n1.1 研究背景\n背景内容"),
                    _FakePdfPage("第二章 方法\n方法内容"),
                ]

        import fileops.document_split as document_split

        monkeypatch.setattr(document_split, "PdfReader", _FakePdfReader)

        out_dir = root / "pdf_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (out_dir / "paper_split_index.json").exists()

        produced_txt = sorted(out_dir.glob("paper_*.txt"))
        assert len(produced_txt) >= 3
        merged_text = "\n".join(path.read_text(encoding="utf-8") for path in produced_txt)
        assert "第一章 绪论" in merged_text
        assert "1.1 研究背景" in merged_text
        assert "第二章 方法" in merged_text


def test_doc_split_pdf_import_format_mismatch() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="markdown",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.FAILED
        assert "import format" in results[0].message.lower()


def test_doc_split_pdf_heading_without_spaces(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "thesis.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        class _FakePdfPage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class _FakePdfReader:
            def __init__(self, _stream: object) -> None:
                self.pages = [
                    _FakePdfPage("封面\n第一章绪论\n正文A\n1.1研究背景\n正文B\n第二章方法\n正文C"),
                ]

        import fileops.document_split as document_split

        monkeypatch.setattr(document_split, "PdfReader", _FakePdfReader)

        out_dir_h1 = root / "pdf_h1"
        h1_results = split_documents_by_structure(
            sources=[src],
            destination=out_dir_h1,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(h1_results) == 1
        assert h1_results[0].status == OperationStatus.SUCCESS
        h1_files = sorted(out_dir_h1.glob("thesis_*.txt"))
        assert len(h1_files) >= 2

        out_dir_h1_h2 = root / "pdf_h1h2"
        h1_h2_results = split_documents_by_structure(
            sources=[src],
            destination=out_dir_h1_h2,
            workspace=root,
            dry_run=False,
            heading_mode="h1_h2",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(h1_h2_results) == 1
        assert h1_h2_results[0].status == OperationStatus.SUCCESS
        h1_h2_files = sorted(out_dir_h1_h2.glob("thesis_*.txt"))
        assert len(h1_h2_files) > len(h1_files)

        h1_h2_text = "\n".join(path.read_text(encoding="utf-8") for path in h1_h2_files)
        assert "第一章绪论" in h1_h2_text
        assert "1.1研究背景" in h1_h2_text
        assert "第二章方法" in h1_h2_text


def test_doc_split_pdf_auto_output_keeps_pdf(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "manual.pdf"

        from pypdf import PdfReader, PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        writer.add_blank_page(width=595, height=842)
        writer.add_blank_page(width=595, height=842)
        with src.open("wb") as stream:
            writer.write(stream)

        import fileops.document_split as document_split

        monkeypatch.setattr(
            document_split,
            "_split_pdf_document",
            lambda *_args, **_kwargs: [
                {"title": "第一章", "lines": ["第一章"], "start_page": 0, "end_page": 1},
                {"title": "第二章", "lines": ["第二章"], "start_page": 1, "end_page": 3},
            ],
        )

        out_dir = root / "pdf_out"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="pdf",
            output_format="auto",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        produced_pdf = sorted(out_dir.glob("manual_*.pdf"))
        assert len(produced_pdf) == 2
        assert not list(out_dir.glob("manual_*.txt"))

        page_counts = [len(PdfReader(str(path)).pages) for path in produced_pdf]
        assert page_counts == [1, 2]


def test_doc_split_non_pdf_export_pdf_not_supported() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "sample.md"
        src.write_text("# A\nBody\n", encoding="utf-8")

        out_dir = root / "doc_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="markdown",
            output_format="pdf",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.FAILED
        assert "only when the input is PDF" in results[0].message


def test_doc_split_pdf_ignores_repeated_running_header(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "thesis.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        class _FakePdfPage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class _FakePdfReader:
            def __init__(self, _stream: object) -> None:
                self.pages = [
                    _FakePdfPage("第一章 绪论\n正文A\n第1页"),
                    _FakePdfPage("第一章 绪论\n正文B\n第2页"),
                    _FakePdfPage("第一章 绪论\n第二章 方法\n正文C\n第3页"),
                    _FakePdfPage("第二章 方法\n正文D\n第4页"),
                ]

        import fileops.document_split as document_split

        monkeypatch.setattr(document_split, "PdfReader", _FakePdfReader)

        out_dir = root / "pdf_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        produced_txt = sorted(out_dir.glob("thesis_*.txt"))
        assert len(produced_txt) == 2

        merged_text = "\n".join(path.read_text(encoding="utf-8") for path in produced_txt)
        assert "正文A" in merged_text
        assert "正文D" in merged_text


def test_doc_split_pdf_prefers_outline_for_chapters(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "thesis.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        class _FakeDest:
            def __init__(self, title: str, page: int) -> None:
                self.title = title
                self.page = page

        class _FakePdfPage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class _FakePdfReader:
            def __init__(self, _stream: object) -> None:
                self.pages = [
                    _FakePdfPage("封面"),
                    _FakePdfPage("摘要"),
                    _FakePdfPage("第一章 绪论 正文"),
                    _FakePdfPage("第二章 相关理论基础 正文"),
                    _FakePdfPage("第三章 模型异构 正文"),
                ]
                self.outline = [
                    _FakeDest("摘要", 1),
                    _FakeDest("第1章 绪论", 2),
                    _FakeDest("第2章 相关理论基础", 3),
                    _FakeDest("第3章 模型异构", 4),
                ]

            def get_destination_page_number(self, node: _FakeDest) -> int:
                return node.page

        import fileops.document_split as document_split

        monkeypatch.setattr(document_split, "PdfReader", _FakePdfReader)

        out_dir = root / "outline_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        produced_txt = sorted(out_dir.glob("thesis_*.txt"))
        assert len(produced_txt) == 5

        index_file = out_dir / "thesis_split_index.json"
        assert index_file.exists()
        index_text = index_file.read_text(encoding="utf-8")
        assert "第1章 绪论" in index_text
        assert "第2章 相关理论基础" in index_text


def test_doc_split_pdf_outline_h2_level(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "thesis.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        class _FakeDest:
            def __init__(self, title: str, page: int) -> None:
                self.title = title
                self.page = page

        class _FakePdfPage:
            def __init__(self, text: str) -> None:
                self._text = text

            def extract_text(self) -> str:
                return self._text

        class _FakePdfReader:
            def __init__(self, _stream: object) -> None:
                self.pages = [
                    _FakePdfPage("封面"),
                    _FakePdfPage("第一章"),
                    _FakePdfPage("1.1 背景"),
                    _FakePdfPage("1.2 方法"),
                    _FakePdfPage("第二章"),
                ]
                self.outline = [
                    _FakeDest("第1章 绪论", 1),
                    [
                        _FakeDest("1.1 研究背景", 2),
                        _FakeDest("1.2 研究方法", 3),
                    ],
                    _FakeDest("第2章 方法", 4),
                ]

            def get_destination_page_number(self, node: _FakeDest) -> int:
                return node.page

        import fileops.document_split as document_split

        monkeypatch.setattr(document_split, "PdfReader", _FakePdfReader)

        out_dir = root / "outline_h2"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h2",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        produced_txt = sorted(out_dir.glob("thesis_*.txt"))
        assert len(produced_txt) == 3
        merged_text = "\n".join(path.read_text(encoding="utf-8") for path in produced_txt)
        assert "1.1 背景" in merged_text
        assert "1.2 方法" in merged_text
        index_text = (out_dir / "thesis_split_index.json").read_text(encoding="utf-8")
        assert "1.1 研究背景" in index_text
        assert "1.2 研究方法" in index_text


def test_doc_split_pdf_encrypted_requires_password(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        src = root / "secret.pdf"
        src.write_bytes(b"%PDF-1.4 fake")

        class _FakePdfReader:
            def __init__(self, _stream: object) -> None:
                self.is_encrypted = True
                self.pages = []
                self.outline = []

            def decrypt(self, _password: str) -> int:
                return 0

        import fileops.document_split as document_split

        monkeypatch.setattr(document_split, "PdfReader", _FakePdfReader)

        out_dir = root / "secret_parts"
        results = split_documents_by_structure(
            sources=[src],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            heading_mode="h1",
            include_image_text=False,
            input_format="pdf",
            output_format="txt",
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.FAILED
        assert "requires a password" in results[0].message
