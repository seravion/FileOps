from __future__ import annotations

import json
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.shared import Pt

import fileops.document_compare as document_compare_module
from fileops.document_compare import analyze_document_with_template, compare_documents_with_template
from fileops.models import OperationStatus


def _build_template(path: Path) -> None:
    doc = Document()
    doc.add_heading("模板标题", level=1)
    body = doc.add_paragraph("模板正文示例")
    body.paragraph_format.first_line_indent = Pt(28)
    body.paragraph_format.line_spacing = 1.5
    doc.add_paragraph("本论文拟分为六章，具体安排如下：")
    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("1.1 研究背景")
    doc.add_heading("参考文献", level=1)
    ref = doc.add_paragraph("[1] Template Reference")
    ref.paragraph_format.left_indent = Pt(21)
    ref.paragraph_format.first_line_indent = Pt(-21)
    doc.add_paragraph("公式示例 (2-1)")
    doc.add_paragraph("图 2-1 示例")
    doc.save(path)


def _build_source_with_mismatches(path: Path) -> None:
    doc = Document()
    title = doc.add_paragraph("1. 绪论")
    title.paragraph_format.first_line_indent = Pt(28)

    para = doc.add_paragraph("正文段落")
    para.paragraph_format.first_line_indent = Pt(20)
    para.paragraph_format.line_spacing = 2.0

    doc.add_paragraph("本论文拟分为六章，具体安排如下：")
    doc.add_paragraph("第1章 绪论")
    doc.add_paragraph("1.1 研究背景")

    doc.add_paragraph("公式内容（2.2）")
    doc.add_paragraph("公式内容（2.1）")
    doc.add_paragraph("图 2.2 示例")
    doc.add_paragraph("图 2.1 示例")

    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 文献A")
    doc.save(path)


def test_document_compare_generates_report_and_adjusted_file() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        template = root / "template.docx"
        source = root / "paper.docx"
        out_dir = root / "out"

        _build_template(template)
        _build_source_with_mismatches(source)

        results = compare_documents_with_template(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (out_dir / "paper_adjusted.docx").exists()
        assert (out_dir / "paper_compare_report.json").exists()
        assert (out_dir / "paper_compare_report.docx").exists()

        report = json.loads((out_dir / "paper_compare_report.json").read_text(encoding="utf-8-sig"))
        assert report["overview"]["template_name"] == template.name
        assert report["summary"]["total_issues"] > 0
        assert report["summary"]["equation_numbering"] > 0
        assert report["summary"]["figure_numbering"] > 0
        assert report["categories"]
        first_issue = report["issues"][0]
        assert "location" in first_issue
        assert "adjustment" in first_issue
        assert "category_label" in first_issue


def test_analyze_document_with_template_returns_summary() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        template = root / "template.docx"
        source = root / "paper.docx"

        _build_template(template)
        _build_source_with_mismatches(source)

        analysis = analyze_document_with_template(source=source, template_path=template)
        assert analysis["overview"]["template_name"] == template.name
        assert analysis["overview"]["source_name"] == source.name
        assert int(analysis["summary"]["total_issues"]) > 0
        assert isinstance(analysis["issues"], list)


def test_document_compare_dry_run_does_not_write_files() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        template = root / "template.docx"
        source = root / "paper.docx"
        out_dir = root / "out"

        _build_template(template)
        _build_source_with_mismatches(source)

        results = compare_documents_with_template(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=True,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.DRY_RUN
        assert not out_dir.exists()


def test_document_compare_supports_custom_text_report_output() -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        template = root / "template.docx"
        source = root / "paper.docx"
        out_dir = root / "out"
        custom_report = root / "compare_detail.txt"

        _build_template(template)
        _build_source_with_mismatches(source)

        results = compare_documents_with_template(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
            detailed_report_path=custom_report,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert custom_report.exists()
        report_text = custom_report.read_text(encoding="utf-8-sig")
        assert "格式检测报告" in report_text
        assert "差异明细" in report_text
        assert "调整建议" in report_text


def test_document_compare_can_generate_ai_assist_file(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        template = root / "template.docx"
        source = root / "paper.docx"
        out_dir = root / "out"

        _build_template(template)
        _build_source_with_mismatches(source)

        def _fake_ai_report(*, analysis: dict[str, object], output_path: Path, config: dict[str, object]) -> Path:
            output_path.write_text("# AI建议\n- 优先修复标题样式\n", encoding="utf-8-sig")
            return output_path

        monkeypatch.setattr(document_compare_module, "generate_compare_ai_report", _fake_ai_report)

        results = compare_documents_with_template(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
            ai_assist_config={
                "api_key": "dummy-key",
                "model": "dummy-model",
                "base_url": "https://example.com/v1",
            },
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS
        assert (out_dir / "paper_ai_assist.md").exists()
        assert "AI建议" in (out_dir / "paper_ai_assist.md").read_text(encoding="utf-8-sig")
