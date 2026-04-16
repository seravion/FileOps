from __future__ import annotations

from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from PIL import Image

from fileops.models import OperationStatus
from fileops.word_template import format_word_documents, import_word_template, list_word_templates


def _create_sample_png(path: Path) -> None:
    image = Image.new("RGB", (24, 24), color=(255, 0, 0))
    image.save(path)


def test_template_library_import_and_list(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        home = Path(temp_dir)
        monkeypatch.setenv("HOME", str(home))
        monkeypatch.setenv("USERPROFILE", str(home))
        monkeypatch.setenv("LOCALAPPDATA", str(home))
        monkeypatch.setenv("APPDATA", str(home))

        template_file = home / "sample_template.docx"
        doc = Document()
        doc.add_paragraph("template")
        doc.save(template_file)

        imported = import_word_template(template_file)
        templates = list_word_templates()

        assert imported.exists()
        assert imported in templates


def test_word_format_applies_template_and_preserves_content(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        monkeypatch.setenv("HOME", str(root))
        monkeypatch.setenv("USERPROFILE", str(root))
        monkeypatch.setenv("LOCALAPPDATA", str(root))
        monkeypatch.setenv("APPDATA", str(root))

        template = root / "template.docx"
        template_doc = Document()
        template_doc.add_heading("Template Title", level=1)
        template_doc.save(template)

        source = root / "source.docx"
        image_path = root / "img.png"
        _create_sample_png(image_path)

        source_doc = Document()
        source_doc.add_heading("Source Heading", level=1)
        source_doc.add_paragraph("Source paragraph")
        table = source_doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "C1"
        table.cell(0, 1).text = "C2"
        source_doc.add_picture(str(image_path))
        source_doc.save(source)

        out_dir = root / "out"
        results = format_word_documents(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        output = out_dir / "source_formatted.docx"
        assert output.exists()

        out_doc = Document(str(output))
        text_blob = "\n".join(par.text for par in out_doc.paragraphs)
        assert "Source Heading" in text_blob
        assert "Source paragraph" in text_blob
        assert any(tbl.cell(0, 0).text == "C1" for tbl in out_doc.tables)
        assert len(out_doc.inline_shapes) >= 1


def test_word_format_dry_run(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        monkeypatch.setenv("HOME", str(root))
        monkeypatch.setenv("USERPROFILE", str(root))
        monkeypatch.setenv("LOCALAPPDATA", str(root))
        monkeypatch.setenv("APPDATA", str(root))

        template = root / "template.docx"
        doc = Document()
        doc.add_paragraph("template")
        doc.save(template)

        source = root / "source.docx"
        source_doc = Document()
        source_doc.add_paragraph("source")
        source_doc.save(source)

        out_dir = root / "out"
        results = format_word_documents(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=True,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.DRY_RUN
        assert not out_dir.exists()


def test_word_format_prefers_template_style_and_clears_direct_format(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        monkeypatch.setenv("HOME", str(root))
        monkeypatch.setenv("USERPROFILE", str(root))
        monkeypatch.setenv("LOCALAPPDATA", str(root))
        monkeypatch.setenv("APPDATA", str(root))

        template = root / "template.docx"
        template_doc = Document()
        normal_style = template_doc.styles["Normal"]
        normal_style.font.size = Pt(14)
        template_doc.add_paragraph("模板正文示例")
        template_doc.save(template)

        source = root / "source.docx"
        source_doc = Document()
        source_doc.add_heading("章节标题", level=1)
        para = source_doc.add_paragraph("需要套用模板的正文")
        para.paragraph_format.first_line_indent = Pt(28)
        para.paragraph_format.line_spacing = 2.0
        run = para.runs[0]
        run.bold = True
        run.italic = True
        source_doc.save(source)

        out_dir = root / "out"
        results = format_word_documents(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        output = out_dir / "source_formatted.docx"
        assert output.exists()

        out_doc = Document(str(output))
        content_paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]
        assert len(content_paragraphs) >= 2

        body_para = content_paragraphs[1]
        assert str(body_para.style.name) == "Normal"
        assert body_para.paragraph_format.first_line_indent is None
        assert body_para.paragraph_format.line_spacing is None
        assert body_para.runs
        assert body_para.runs[0].bold is None
        assert body_para.runs[0].italic is None


def test_word_format_keeps_toc_level_styles_from_template(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        monkeypatch.setenv("HOME", str(root))
        monkeypatch.setenv("USERPROFILE", str(root))
        monkeypatch.setenv("LOCALAPPDATA", str(root))
        monkeypatch.setenv("APPDATA", str(root))

        template = root / "template.docx"
        template_doc = Document()
        if "目录 1" not in [style.name for style in template_doc.styles]:
            template_doc.styles.add_style("目录 1", WD_STYLE_TYPE.PARAGRAPH)
        if "目录 2" not in [style.name for style in template_doc.styles]:
            template_doc.styles.add_style("目录 2", WD_STYLE_TYPE.PARAGRAPH)
        template_doc.add_paragraph("目录", style="目录 1")
        template_doc.add_paragraph("1 绪论..................1", style="目录 1")
        template_doc.add_paragraph("1.1 研究背景............2", style="目录 2")
        template_doc.save(template)

        source = root / "source.docx"
        source_doc = Document()
        if "目录 1" not in [style.name for style in source_doc.styles]:
            source_doc.styles.add_style("目录 1", WD_STYLE_TYPE.PARAGRAPH)
        if "目录 2" not in [style.name for style in source_doc.styles]:
            source_doc.styles.add_style("目录 2", WD_STYLE_TYPE.PARAGRAPH)
        source_doc.add_paragraph("目录", style="目录 1")
        source_doc.add_paragraph("1 绪论..................1", style="目录 1")
        source_doc.add_paragraph("1.1 研究背景............2", style="目录 2")
        source_doc.save(source)

        out_dir = root / "out"
        results = format_word_documents(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        output = out_dir / "source_formatted.docx"
        out_doc = Document(str(output))
        paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]
        assert len(paragraphs) >= 3
        assert str(paragraphs[1].style.name) == "目录 1"
        assert str(paragraphs[2].style.name) == "目录 2"


def test_word_format_manual_toc_block_uses_template_toc_styles(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        monkeypatch.setenv("HOME", str(root))
        monkeypatch.setenv("USERPROFILE", str(root))
        monkeypatch.setenv("LOCALAPPDATA", str(root))
        monkeypatch.setenv("APPDATA", str(root))

        template = root / "template.docx"
        template_doc = Document()
        if "目录 1" not in [style.name for style in template_doc.styles]:
            template_doc.styles.add_style("目录 1", WD_STYLE_TYPE.PARAGRAPH)
        if "目录 2" not in [style.name for style in template_doc.styles]:
            template_doc.styles.add_style("目录 2", WD_STYLE_TYPE.PARAGRAPH)
        template_doc.add_paragraph("本论文拟分为六章，具体安排如下：")
        template_doc.add_paragraph("第1章 绪论", style="目录 1")
        template_doc.add_paragraph("1.1 研究背景", style="目录 2")
        template_doc.save(template)

        source = root / "source.docx"
        source_doc = Document()
        source_doc.add_paragraph("本论文拟分为六章，具体安排如下：")
        source_doc.add_paragraph("第1章 绪论")
        source_doc.add_paragraph("1.1 研究背景")
        source_doc.save(source)

        out_dir = root / "out"
        results = format_word_documents(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        output = out_dir / "source_formatted.docx"
        out_doc = Document(str(output))
        paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]
        assert len(paragraphs) >= 3
        assert str(paragraphs[1].style.name) == "目录 1"
        assert str(paragraphs[2].style.name) == "目录 2"


def test_word_format_reference_entry_uses_template_hanging_indent(monkeypatch) -> None:
    with TemporaryDirectory() as temp_dir:
        root = Path(temp_dir)
        monkeypatch.setenv("HOME", str(root))
        monkeypatch.setenv("USERPROFILE", str(root))
        monkeypatch.setenv("LOCALAPPDATA", str(root))
        monkeypatch.setenv("APPDATA", str(root))

        template = root / "template.docx"
        template_doc = Document()
        template_doc.add_heading("参考文献", level=1)
        ref = template_doc.add_paragraph("[1] Template Ref")
        ref.paragraph_format.left_indent = Pt(21)
        ref.paragraph_format.first_line_indent = Pt(-21)
        template_doc.save(template)

        source = root / "source.docx"
        source_doc = Document()
        source_doc.add_paragraph("参考文献")
        source_doc.add_paragraph("[1] Source Ref")
        source_doc.save(source)

        out_dir = root / "out"
        results = format_word_documents(
            sources=[source],
            destination=out_dir,
            workspace=root,
            dry_run=False,
            template_path=template,
        )

        assert len(results) == 1
        assert results[0].status == OperationStatus.SUCCESS

        output = out_dir / "source_formatted.docx"
        out_doc = Document(str(output))
        paragraphs = [p for p in out_doc.paragraphs if p.text.strip()]
        assert len(paragraphs) >= 2
        ref_para = paragraphs[1]
        assert str(ref_para.style.name) == "Normal"
        assert getattr(ref_para.paragraph_format.left_indent, "twips", None) == 420
        assert getattr(ref_para.paragraph_format.first_line_indent, "twips", None) == -420
